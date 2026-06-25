import base64
import json
import tempfile
import unittest
import zipfile
from pathlib import Path

from docx import Document

from app.workflow import build_package, generate_documents, generate_source_document
from app.project_generator import generate_java_project


class MaterialDocumentTests(unittest.TestCase):
    PNG_BYTES = base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVQIHWP4z8DwHwAFgAI/ScL6kQAAAABJRU5ErkJggg=="
    )

    def test_documents_include_real_screenshots_and_operation_steps(self):
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            shots = root / "screenshots"
            shots.mkdir()
            planning = {
                "software_name": "测试管理系统",
                "description": "用于测试图文材料生成。",
                "software_type": "管理软件",
                "target_users": "业务人员",
                "modules": [{"key": "items", "name": "信息管理", "description": "管理业务信息", "pages": ["信息列表"], "fields": ["名称", "状态"]}],
                "database_tables": ["business_items"],
                "api_list": ["GET /api/items", "POST /api/items", "PUT /api/items/{id}/approve"],
            }
            (root / "planning.json").write_text(json.dumps(planning, ensure_ascii=False), encoding="utf-8")
            (root / "code_stats.json").write_text('{"total_lines": 200}', encoding="utf-8")
            files = ["01-login.png", "02-dashboard.png", "03-信息管理-list.png", "04-信息管理-create.png"]
            for filename in files:
                (shots / filename).write_bytes(self.PNG_BYTES)
            (root / "screenshot_manifest.json").write_text(json.dumps({"screenshots": [
                {"kind": "login", "file": files[0]}, {"kind": "dashboard", "file": files[1]},
                {"kind": "module_list", "file": files[2], "module_key": "items", "module_name": "信息管理"},
                {"kind": "module_create", "file": files[3], "module_key": "items", "module_name": "信息管理"},
            ]}, ensure_ascii=False), encoding="utf-8")

            generate_documents({}, root)

            design = Document(root / "docs" / "设计说明书.docx")
            manual = Document(root / "docs" / "用户操作手册.docx")
            from PIL import Image
            with Image.open(root / "design_architecture.png") as image:
                self.assertEqual(image.size, (1800, 900))
            with Image.open(root / "design_process.png") as image:
                self.assertEqual(image.size, (1800, 700))
            self.assertGreaterEqual(len(design.inline_shapes), 5)
            self.assertGreaterEqual(len(manual.inline_shapes), 4)
            design_text = "\n".join(paragraph.text for paragraph in design.paragraphs)
            self.assertIn("分层系统架构图", design_text)
            self.assertIn("系统业务处理流程图", design_text)
            self.assertIn("功能目标", design_text)
            manual_text = "\n".join(paragraph.text for paragraph in manual.paragraphs)
            self.assertIn("9.1.2 操作说明", manual_text)
            self.assertIn("目录", manual_text)
            self.assertIn("产品背景与目标", manual_text)
            self.assertGreaterEqual(len(manual.inline_shapes), 5)
            self.assertIn("第1步：", manual_text)
            self.assertEqual(len(manual.tables), 1)
            self.assertIn("第 1 页", "\n".join(paragraph.text for paragraph in manual.sections[0].header.paragraphs))

    def test_source_material_excludes_ai_markers_and_compressed_css(self):
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            project = root / "generated_project"
            (project / "frontend/src").mkdir(parents=True)
            (project / "backend/src/main/java/example").mkdir(parents=True)
            (root / "planning.json").write_text(
                json.dumps({"software_name": "测试管理系统"}, ensure_ascii=False), encoding="utf-8"
            )
            (project / "backend/src/main/java/example/ItemsService.java").write_text(
                "class ItemsService { void save() {} }", encoding="utf-8"
            )
            (project / "frontend/src/style.css").write_text(
                "/* AI UI Enhancer: shell */ body{color:red}", encoding="utf-8"
            )

            generate_source_document(root)

            source = Document(root / "docs" / "源代码材料.docx")
            text = "\n".join(paragraph.text for paragraph in source.paragraphs).lower()
            self.assertNotIn("ai ui enhancer", text)
            self.assertNotIn("style.css", text)
            self.assertIn("itemsservice", text)
            self.assertIn("源代码", "\n".join(paragraph.text for paragraph in source.sections[0].header.paragraphs))
            self.assertIn("1/1", "\n".join(paragraph.text for paragraph in source.sections[0].footer.paragraphs))
            with zipfile.ZipFile(root / "docs" / "源代码材料.docx") as archive:
                document_xml = archive.read("word/document.xml").decode("utf-8")
            self.assertNotIn('w:type="page"', document_xml)

    def test_generated_view_exposes_stable_screenshot_action_selectors(self):
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            planning = {
                "software_name": "测试管理系统",
                "description": "用于测试生成页面。",
                "software_type": "管理软件",
                "industry_name": "教育",
                "modules": [{"key": "items", "name": "信息管理", "description": "管理业务信息", "pages": ["信息列表"], "fields": ["名称", "状态"]}],
                "database_tables": ["business_items"],
                "api_list": ["GET /api/items", "POST /api/items", "PUT /api/items/{id}/approve"],
                "ui_plan": {"shell": "sidebar_admin", "home_pattern": "metric_dashboard", "module_patterns": {"items": "table_crud"}},
            }
            (root / "planning.json").write_text(json.dumps(planning, ensure_ascii=False), encoding="utf-8")
            generate_java_project(root)
            page = (root / "generated_project/frontend/src/views/ItemsPage.vue").read_text(encoding="utf-8")
            self.assertIn('data-action="create"', page)
            self.assertIn('data-action="edit"', page)
            self.assertIn('data-action="delete"', page)
            self.assertIn(':data-action="\'approve\'"', page)

    def test_final_package_contains_screenshot_manifest(self):
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            (root / "generated_project").mkdir()
            (root / "generated_project/README.md").write_text("generated", encoding="utf-8")
            (root / "docs").mkdir()
            (root / "screenshots").mkdir()
            (root / "planning.json").write_text(json.dumps({"software_name": "测试管理系统"}, ensure_ascii=False), encoding="utf-8")
            (root / "code_stats.json").write_text(json.dumps({"total_lines": 1, "frontend_lines": 1, "backend_lines": 0, "sql_lines": 0}), encoding="utf-8")
            (root / "screenshot_manifest.json").write_text('{"screenshots": []}', encoding="utf-8")
            (root / "document_narratives.json").write_text('{"meta": {"mode": "template"}, "modules": {}}', encoding="utf-8")

            build_package(root)

            with zipfile.ZipFile(root / "copyright_package.zip") as package:
                self.assertIn("screenshot_manifest.json", package.namelist())
                self.assertIn("document_narratives.json", package.namelist())


if __name__ == "__main__":
    unittest.main()
