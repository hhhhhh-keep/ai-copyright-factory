import json
import logging
import os
import re
import shutil
import socket
import subprocess
import sys
import threading
import time
import textwrap
import uuid
import zipfile
from datetime import datetime, timedelta
from multiprocessing import Process
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from .compliance import build_compliance_report
from .document_narratives import build_document_narratives
from .enhancer import ALLOWED_FILES, enhance_project, restore_enhancement
from .planner import PlannerValidationError, build_planning
from .project_generator import generate_java_project


BASE_DIR = Path(__file__).resolve().parents[2]
OUTPUT_ROOT = BASE_DIR / "outputs"
OUTPUT_ROOT.mkdir(parents=True, exist_ok=True)

# ISSUE-026：补模块级 logger。ISSUE-023 P1-3 在 continue_material_generation 末尾
# 新增 logger.warning 调用但未在模块顶部 import，触发 NameError 阻断软著材料打包。
logger = logging.getLogger(__name__)
_LOCK = threading.Lock()

# ISSUE-008 L1 状态与超时阈值：超过这个时间且进程不存在/锁文件陈旧
# 的任务会被启动扫描标记为 interrupted。
INTERRUPTED_STATUS = "interrupted"
EXECUTING_STATUSES = {
    "generating",
    "regenerating_project",
    "generating_materials",
    "confirmed",
    "awaiting_demo_review",
}
# 不同状态下"updated_at" 超过这个秒数即视为陈旧
INTERRUPT_AGE_SECONDS = {
    "confirmed": 600,  # 10 分钟，规划已确认但项目目录还没建
    "generating": 1800,  # 30 分钟，生成项目阶段
    "regenerating_project": 1800,
    "generating_materials": 3600,  # 60 分钟，材料阶段
}
# awaiting_demo_review 不自动标 interrupted（可能用户长时间不操作）

STEPS = [
    ("planning", "生成软件规划"),
    ("project", "生成可运行项目"),
    ("enhance", "AI 增强项目代码"),
    ("run", "运行项目并验证"),
    ("demo", "启动在线 Demo"),
    ("screenshot", "自动截图"),
    ("analyze", "统计真实代码行数"),
    ("source", "生成源码材料"),
    ("docs", "生成设计说明书和用户手册"),
    ("compliance", "执行软著合规检查"),
    ("package", "打包软著材料"),
]


def _now() -> str:
    return datetime.now().isoformat(timespec="seconds")


def _json_write(path: Path, data: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    temporary = path.with_suffix(path.suffix + f".{os.getpid()}.tmp")
    temporary.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
    os.replace(str(temporary), str(path))


def _json_read(path: Path) -> Dict[str, Any]:
    last_error: Optional[OSError] = None
    for attempt in range(5):
        try:
            return json.loads(path.read_text(encoding="utf-8-sig"))
        except PermissionError as exc:
            # Windows may briefly deny readers while os.replace/write handles
            # are settling. Status polling should wait instead of surfacing a
            # transient 500 to the frontend.
            last_error = exc
            time.sleep(0.05 * (attempt + 1))
    if last_error:
        raise last_error
    return json.loads(path.read_text(encoding="utf-8-sig"))


def _set_demo_stage(job_id: str, **changes: Any) -> None:
    """更新 demo_runtime.json 的 stage / error / 进度等字段。

    不修改 status="running" 的运行态；status 仍由 start_online_demo 收尾时统一写。
    """
    path = OUTPUT_ROOT / job_id / "demo_runtime.json"
    with _LOCK:
        if not path.exists():
            runtime = {
                "status": "starting",
                "stage": "queued",
                "started_at": _now(),
                "expires_at": (datetime.now() + timedelta(hours=1)).isoformat(
                    timespec="seconds"
                ),
            }
        else:
            runtime = _json_read(path)
        runtime.update(changes)
        runtime["updated_at"] = _now()
        _json_write(path, runtime)


def _status_path(job_id: str) -> Path:
    return OUTPUT_ROOT / job_id / "status.json"


# ----------------- ISSUE-008 L1：worker 锁与中断检测 -----------------


def _worker_lock_path(job_id: str) -> Path:
    return OUTPUT_ROOT / job_id / "worker.lock"


def _pid_alive(pid: Optional[int]) -> bool:
    """检查 PID 是否仍存在。Windows 上 os.kill(pid, 0) 抛 OSError 即死亡。"""
    if not pid or pid <= 0:
        return False
    try:
        os.kill(pid, 0)
        return True
    except (OSError, ProcessLookupError):
        return False


def _acquire_worker_lock(job_id: str, task_name: str) -> Dict[str, Any]:
    """把当前 Worker 写入 outputs/{job_id}/worker.lock。

    返回 lock 字典（含 pid、started_at、task）。如果 lock 已存在且 PID 仍存活，
    抛 RuntimeError，让调用方拒绝重复启动。
    """
    job_dir = OUTPUT_ROOT / job_id
    job_dir.mkdir(parents=True, exist_ok=True)
    lock_path = _worker_lock_path(job_id)
    if lock_path.exists():
        try:
            existing = _json_read(lock_path)
        except (OSError, ValueError):
            existing = {}
        existing_pid = existing.get("pid")
        if _pid_alive(existing_pid):
            raise RuntimeError(
                f"任务 {job_id} 已有活跃 Worker (PID {existing_pid})，拒绝重复启动"
            )
    lock = {
        "pid": os.getpid(),
        "ppid": os.getppid(),
        "task": task_name,
        "started_at": _now(),
    }
    _json_write(lock_path, lock)
    return lock


def _release_worker_lock(job_id: str) -> None:
    """删除 worker.lock。"""
    lock_path = _worker_lock_path(job_id)
    if lock_path.exists():
        try:
            lock_path.unlink()
        except OSError:
            pass


def _is_job_orphaned(status: Dict[str, Any]) -> bool:
    """根据 lock 文件和 updated_at 判断任务是否需要标 interrupted。"""
    job_id = status.get("job_id")
    if not job_id:
        return False
    if status.get("status") not in EXECUTING_STATUSES:
        return False
    # success / failed / interrupted / draft_planning 直接跳过
    if status.get("status") in {"success", "failed", INTERRUPTED_STATUS, "draft_planning"}:
        return False
    # awaiting_demo_review：长时间没动可能用户没操作，但应给机会先看 UI
    # 暂时保守：awaiting_demo_review 不自动 interrupted（用户可手动恢复）
    if status.get("status") == "awaiting_demo_review":
        return False
    # 优先级 1：worker.lock 是否存在
    # - 活锁：Worker 仍在跑，即使 updated_at 旧也不能算孤儿（避免长任务被误杀）
    # - 死锁：Worker 已退出，视为孤儿
    # - 坏锁：lock JSON 解析失败，视为孤儿
    lock_path = _worker_lock_path(job_id)
    if lock_path.exists():
        try:
            lock = _json_read(lock_path)
        except (OSError, ValueError):
            return True
        if not _pid_alive(lock.get("pid")):
            return True
        return False
    # 优先级 2：无锁 + updated_at 超时 → 视为孤儿
    # （任务曾经启动过 worker，但 worker 没写 lock 就崩溃了，靠 updated_at 兜底）
    updated_at = status.get("updated_at")
    if updated_at:
        try:
            last = datetime.fromisoformat(updated_at)
            age = (datetime.now() - last).total_seconds()
            threshold = INTERRUPT_AGE_SECONDS.get(status["status"], 1800)
            if age >= threshold:
                return True
        except ValueError:
            pass
    return False


def scan_for_interrupted_jobs() -> List[str]:
    """扫描 outputs/*/status.json，对疑似失联的任务标记 interrupted。

    返回被标记的 job_id 列表。服务启动时调用一次。
    """
    marked: List[str] = []
    if not OUTPUT_ROOT.exists():
        return marked
    for status_path in OUTPUT_ROOT.glob("*/status.json"):
        try:
            status = _json_read(status_path)
        except (OSError, ValueError):
            continue
        job_id = status.get("job_id")
        if not job_id:
            continue
        if not _is_job_orphaned(status):
            continue
        try:
            # 标记前先停掉任何残留 Demo
            try:
                stop_online_demo(job_id)
            except Exception:
                pass
            steps = status.get("steps", [])
            for item in steps:
                if item.get("status") == "running":
                    item["status"] = "failed"
            _update(
                job_id,
                status=INTERRUPTED_STATUS,
                steps=steps,
                current_step="后台进程中断，等待用户恢复",
                failed_stage=status.get("status"),
                interrupted_at=_now(),
                interrupted_reason=(
                    "updated_at 超过阈值或 worker 锁进程已退出，"
                    "请在历史任务页点击'恢复任务'继续。"
                ),
                error=None,
            )
            _release_worker_lock(job_id)
            marked.append(job_id)
        except Exception:
            # 单个任务标记失败不影响其它任务
            continue
    return marked


def _has_active_worker_lock(job_id: str) -> bool:
    """只读检查：lock 文件存在且 PID 仍存活。

    与 `_acquire_worker_lock` 不同，这里不抢锁，只用于 resume 前防双跑。
    """
    lock_path = _worker_lock_path(job_id)
    if not lock_path.exists():
        return False
    try:
        lock = _json_read(lock_path)
    except (OSError, ValueError):
        return False
    return _pid_alive(lock.get("pid"))


def resume_job(job_id: str) -> Dict[str, Any]:
    """根据任务当前阶段选择恢复点，启动新 Worker。

    恢复策略（与 ISSUE-008 文档表一致）：
    - confirmed：项目目录不存在 → 从 project 开始
    - project/enhance：清掉 generated_project/ 后从 project 开始
    - run：从 run 重新执行
    - demo：清掉残留进程后从 demo 重新启动
    - generating_materials：清掉 screenshots/ docs/ copyright_package.zip 等，从 materials 整体重做
    - 其余状态（含 success / draft_planning / failed / awaiting_demo_review）：
      抛 ValueError，由 main.py 转为 409
    - 防双跑：若 lock 文件中 PID 仍存活则拒绝，避免重复点击启动两个生成进程。

    返回 dict：被恢复任务的最新 status。
    """
    status = get_job(job_id)
    if not status:
        raise LookupError("任务不存在")
    current = status.get("status")
    if current == "success":
        raise ValueError("任务已完成，无需恢复")
    if current == "draft_planning":
        raise ValueError("规划未确认，无需恢复")
    if current == "awaiting_demo_review":
        raise ValueError("任务正在等待用户审查 Demo，无需恢复")
    if current != INTERRUPTED_STATUS:
        raise ValueError(f"当前状态 {current} 不可恢复，请先中断或失败后重试")

    # 防双跑：先释放陈旧 lock（PID 已死），再检查是否还有活锁
    lock_path = _worker_lock_path(job_id)
    if lock_path.exists():
        try:
            existing = _json_read(lock_path)
        except (OSError, ValueError):
            existing = {}
        if not _pid_alive(existing.get("pid")):
            _release_worker_lock(job_id)
    if _has_active_worker_lock(job_id):
        raise RuntimeError("已有活跃 Worker 在执行该任务，请稍候再试")

    job_dir = OUTPUT_ROOT / job_id
    # 清理可能残留的运行中 step
    steps = status.get("steps", [])
    for item in steps:
        if item["key"] == "planning":
            item["status"] = "completed"
        elif item["status"] == "running":
            item["status"] = "failed"

    # 根据中断前的状态选择恢复点
    prev_status = status.get("failed_stage") or current
    target = "project"  # 默认从 project 开始
    if prev_status in {"generating_materials", "materials"}:
        # 材料阶段：内部 target 落到第一个材料步骤 screenshot，
        # 避免 STEPS 中没有"materials"导致从 planning 全部重置
        target = "screenshot"
        for name in (
            "screenshots",
            "docs",
            "logs",
            "copyright_package.zip",
            "generated_project.zip",
        ):
            path = job_dir / name
            if path.is_dir():
                shutil.rmtree(path)
            elif path.exists():
                path.unlink()
    elif prev_status == "confirmed":
        target = "project"
    elif prev_status in {"project", "enhance"}:
        target = "project"
        gen_dir = job_dir / "generated_project"
        if gen_dir.exists():
            shutil.rmtree(gen_dir)
    elif prev_status == "run":
        target = "run"
    elif prev_status == "demo":
        target = "demo"
        try:
            stop_online_demo(job_id)
        except Exception:
            pass

    # 解析 target 在 STEPS 中的位置；若 target 不在 STEPS，target=project 时回退到 1 (project)
    # materials 已在前面改写为 screenshot（STEPS 中存在），不会再走这里。
    target_idx = next(
        (i for i, s in enumerate(STEPS) if s[0] == target),
        1 if target in {"project", "enhance"} else 0,
    )

    # 重置 target 之后（含）所有 step 状态为 pending
    for i, (key, _name) in enumerate(STEPS):
        if i >= target_idx:
            for item in steps:
                if item["key"] == key:
                    item["status"] = "pending"

    resume_count = int(status.get("resume_count", 0)) + 1
    # 选择目标状态与 Worker：材料阶段直接进入 generating_materials，避免短暂写 regenerating_project
    if target == "screenshot" or target == "materials":
        new_status = "generating_materials"
        target_fn = continue_material_generation
    else:
        new_status = "generating"
        target_fn = run_job
    _update(
        job_id,
        status=new_status,
        current_step=f"恢复任务，从 {target} 重新开始",
        steps=steps,
        error=None,
        failed_stage=None,
        interrupted_at=None,
        interrupted_reason=None,
        recovery_from_step=target,
        resume_count=resume_count,
        resumed_at=_now(),
    )

    Process(
        target=target_fn,
        args=(job_id,),
        name=f"resume-{target}-{job_id}",
        daemon=True,
    ).start()
    return get_job(job_id)


def create_job(payload: Dict[str, Any]) -> Dict[str, Any]:
    job_id = datetime.now().strftime("%Y%m%d%H%M%S") + "-" + uuid.uuid4().hex[:8]
    job_dir = OUTPUT_ROOT / job_id
    job_dir.mkdir(parents=True)
    status = {
        "job_id": job_id,
        "software_name": payload["software_name"].strip(),
        "description": payload.get("description", "").strip(),
        "software_type": payload.get("software_type", "管理系统").strip(),
        "industry_type": payload["industry_type"],
        "clarification_answers": payload.get("clarification_answers", {}),
        "planner_model": None,
        "codegen_mode": payload.get("codegen_mode", "auto"),
        "codegen_actual_mode": None,
        "codegen_model": None,
        "codegen_summary": None,
        "codegen_changed_files": [],
        "codegen_fallback_reason": None,
        "codegen_enhance_steps": [],
        "document_template": payload.get("document_template", "standard"),
        "version": payload.get("version", "V1.0"),
        "applicant_name": payload.get("applicant_name", "待填写"),
        "completion_date": payload.get("completion_date")
        or datetime.now().strftime("%Y-%m-%d"),
        "publication_status": payload.get("publication_status", "未发表"),
        "compliance_score": None,
        "compliance_grade": None,
        "compliance_passed": None,
        "demo_url": None,
        "swagger_url": None,
        "run_status": "pending",
        "run_validation": None,
        "status": "generating",
        "progress": 0,
        "current_step": "生成软件规划",
        "steps": [{"key": key, "name": name, "status": "pending"} for key, name in STEPS],
        "failed_stage": None,
        "error": None,
        # ISSUE-008 L1：执行信息
        "worker_pid": None,
        "worker_started_at": None,
        "worker_heartbeat_at": None,
        "interrupted_at": None,
        "interrupted_reason": None,
        "resume_count": 0,
        "recovery_from_step": None,
        "resumed_at": None,
        "created_at": _now(),
        "updated_at": _now(),
    }
    _json_write(_status_path(job_id), status)
    return status


def get_job(job_id: str) -> Optional[Dict[str, Any]]:
    path = _status_path(job_id)
    if not path.exists():
        return None
    return _json_read(path)


def _update(job_id: str, **changes: Any) -> Dict[str, Any]:
    with _LOCK:
        status = _json_read(_status_path(job_id))
        status.update(changes)
        status["updated_at"] = _now()
        _json_write(_status_path(job_id), status)
        return status


def _step(job_id: str, key: str, state: str) -> None:
    status = get_job(job_id)
    if not status:
        return
    for item in status["steps"]:
        if item["key"] == key:
            item["status"] = state
    completed = len([item for item in status["steps"] if item["status"] == "completed"])
    current = next((name for step_key, name in STEPS if step_key == key), key)
    _update(
        job_id,
        steps=status["steps"],
        current_step=current,
        progress=int(completed / len(STEPS) * 100),
    )


def generate_planning_draft(job_id: str) -> None:
    job = get_job(job_id)
    if not job:
        return
    job_dir = OUTPUT_ROOT / job_id
    try:
        _acquire_worker_lock(job_id, "generate_planning_draft")
    except RuntimeError:
        # 已有活跃 Worker，跳过启动
        return
    try:
        _update(
            job_id,
            status="generating",
            current_step="生成软件规划",
            error=None,
            failed_stage=None,
            worker_pid=os.getpid(),
            worker_started_at=_now(),
            worker_heartbeat_at=_now(),
        )
        _step(job_id, "planning", "running")
        generate_planning(job, job_dir)
        _step(job_id, "planning", "completed")
        _update(
            job_id,
            status="draft_planning",
            progress=10,
            current_step="等待确认软件规划",
            failed_stage=None,
        )
    except Exception as exc:
        _write_planner_diagnostics(job_dir, exc)
        status = get_job(job_id)
        if status:
            for item in status["steps"]:
                if item["status"] == "running":
                    item["status"] = "failed"
            _update(
                job_id,
                status="failed",
                steps=status["steps"],
                current_step="规划生成失败",
                failed_stage="planning",
                error=f"{type(exc).__name__}: {exc}",
            )
    finally:
        _release_worker_lock(job_id)


def run_job(job_id: str) -> None:
    job = get_job(job_id)
    if not job:
        return
    job_dir = OUTPUT_ROOT / job_id
    if not (job_dir / "planning.json").exists():
        _update(
            job_id,
            status="failed",
            current_step="生成失败",
            failed_stage="project",
            error="planning.json 不存在，禁止开始项目生成",
        )
        return
    try:
        _acquire_worker_lock(job_id, "run_job")
    except RuntimeError:
        return
    try:
        _update(
            job_id,
            status="generating",
            current_step="开始生成",
            error=None,
            failed_stage=None,
            worker_pid=os.getpid(),
            worker_started_at=_now(),
            worker_heartbeat_at=_now(),
        )
        tasks = [
            ("project", lambda: generate_project(job_dir)),
            ("enhance", lambda: enhance_generated_project(job, job_dir)),
            ("run", lambda: validate_generated_project(job, job_dir)),
            ("demo", lambda: start_online_demo(job, job_dir)),
        ]
        for key, action in tasks:
            _step(job_id, key, "running")
            action()
            _step(job_id, key, "completed")
        _update(
            job_id,
            status="awaiting_demo_review",
            progress=int(5 / len(STEPS) * 100),
            current_step="等待用户审查在线 Demo",
            review_round=int(job.get("review_round", 0)) + 1,
        )
    except Exception as exc:
        status = get_job(job_id)
        if status:
            for item in status["steps"]:
                if item["status"] == "running":
                    item["status"] = "failed"
            _update(
                job_id,
                status="failed",
                steps=status["steps"],
                current_step="生成失败",
                failed_stage="project",
                error=f"{type(exc).__name__}: {exc}",
            )
    finally:
        _release_worker_lock(job_id)


def continue_material_generation(job_id: str) -> None:
    job = get_job(job_id)
    if not job:
        return
    job_dir = OUTPUT_ROOT / job_id
    try:
        _acquire_worker_lock(job_id, "continue_material_generation")
    except RuntimeError:
        return
    try:
        _update(
            job_id,
            status="generating_materials",
            current_step="Demo 审查通过，开始生成软著材料",
            error=None,
            failed_stage=None,
            worker_pid=os.getpid(),
            worker_started_at=_now(),
            worker_heartbeat_at=_now(),
        )
        tasks = [
            ("screenshot", lambda: capture_screenshots(job_dir)),
            ("analyze", lambda: analyze_code(job_dir)),
            ("source", lambda: generate_source_document(job_dir, job)),
            ("docs", lambda: generate_documents(job, job_dir)),
            ("compliance", lambda: run_compliance_check(job, job_dir)),
            ("package", lambda: build_package(job_dir)),
        ]
        for key, action in tasks:
            _step(job_id, key, "running")
            action()
            _step(job_id, key, "completed")
        _update(job_id, status="success", progress=100, current_step="生成完成")
        # ISSUE-022：材料生成完成时若 run_status 仍是 running，说明 demo 进程没显式停止。
        # 仅打 warning，不主动 kill（用户可能正在浏览器看 demo）。
        # 完整 daemon 调度改造留 ISSUE-008 L2。
        final = get_job(job_id)
        if final and final.get("run_status") == "running":
            logger.warning(
                "jobId=%s 材料生成已完成，但 run_status=running，demo 进程未显式停止；"
                "建议前端提示用户『停止 Demo』或调整过期清理逻辑。",
                job_id,
            )
    except Exception as exc:
        status = get_job(job_id)
        if status:
            for item in status["steps"]:
                if item["status"] == "running":
                    item["status"] = "failed"
            _update(
                job_id,
                status="failed",
                steps=status["steps"],
                current_step="材料生成失败",
                failed_stage="materials",
                error=f"{type(exc).__name__}: {exc}",
            )
    finally:
        _release_worker_lock(job_id)


def save_planning_version(
    job_id: str,
    planning: Dict[str, Any],
    instruction: str = "",
    summary: str = "",
) -> int:
    versions_dir = OUTPUT_ROOT / job_id / "planning_versions"
    versions_dir.mkdir(parents=True, exist_ok=True)
    version = len(list(versions_dir.glob("v*.json"))) + 1
    _json_write(
        versions_dir / f"v{version}.json",
        {
            "version": version,
            "created_at": _now(),
            "instruction": instruction,
            "summary": summary,
            "planning": planning,
        },
    )
    return version


def reset_job_for_revision(job_id: str) -> Dict[str, Any]:
    job = get_job(job_id)
    if not job:
        raise ValueError("任务不存在")
    stop_online_demo(job_id)
    job_dir = OUTPUT_ROOT / job_id
    for name in (
        "generated_project",
        ".enhancer_backup",
        "screenshots",
        "docs",
        "logs",
        "copyright_package.zip",
        "generated_project.zip",
        "enhancement.json",
        "run_validation.json",
        "code_stats.json",
        "compliance_report.json",
        "demo_runtime.json",
    ):
        path = job_dir / name
        if path.is_dir():
            shutil.rmtree(path)
        elif path.exists():
            path.unlink()
    for item in job["steps"]:
        if item["key"] != "planning":
            item["status"] = "pending"
    return _update(
        job_id,
        status="regenerating_project",
        progress=10,
        current_step="按修改后的规划重新生成项目",
        steps=job["steps"],
        run_status="pending",
        demo_url=None,
        swagger_url=None,
        error=None,
    )


def generate_planning(job: Dict[str, Any], job_dir: Path) -> None:
    result = build_planning(job)
    planning = result.planning.model_dump()
    planning["planner"] = {
        "model": result.model,
    }
    _json_write(job_dir / "planning.json", planning)
    _update(
        job["job_id"],
        planner_model=result.model,
    )


def _write_planner_diagnostics(job_dir: Path, exc: Exception) -> None:
    if not isinstance(exc, PlannerValidationError):
        return
    diagnostics_dir = job_dir / "planner_diagnostics"
    diagnostics_dir.mkdir(parents=True, exist_ok=True)
    (diagnostics_dir / "planner_raw_initial.txt").write_text(
        exc.first_text or "",
        encoding="utf-8",
    )
    (diagnostics_dir / "planner_raw_repair.txt").write_text(
        exc.second_text or "",
        encoding="utf-8",
    )
    _json_write(
        diagnostics_dir / "planner_diagnostics.json",
        {
            "created_at": _now(),
            "error": str(exc),
            "first_error": exc.first_error,
            "second_error": exc.second_error,
            "initial_response_file": "planner_raw_initial.txt",
            "repair_response_file": "planner_raw_repair.txt",
        },
    )


def enhance_generated_project(job: Dict[str, Any], job_dir: Path) -> None:
    planning = _json_read(job_dir / "planning.json")
    requested_mode = (job.get("codegen_mode") or "auto").strip().lower()

    from .enhancer import UI_ENHANCEMENT_STEPS

    # ISSUE-020：UI 增强拆成 5 个子步骤（theme 仅出令牌不写文件）+ README 文档
    UI_STEP_KEYS: Tuple[str, ...] = tuple(key for key, _ in UI_ENHANCEMENT_STEPS)
    UI_STEP_NAMES: Dict[str, str] = dict(UI_ENHANCEMENT_STEPS)
    README_KEY = "readme"

    enhance_steps: List[Dict[str, Any]] = []
    if requested_mode != "template":
        # 顺序：theme → shell → business → dashboard → responsive → readme
        all_keys = list(UI_STEP_KEYS) + [README_KEY]
        all_names = {**UI_STEP_NAMES, README_KEY: "项目说明"}
        enhance_steps = [
            {
                "key": key,
                "file": (
                    "frontend/src/style.css"
                    if key in UI_STEP_KEYS
                    else "README.md"
                ),
                "name": all_names[key],
                "kind": "ui" if key in UI_STEP_KEYS else "doc",
                "status": "pending",
                "summary": "",
                "attempts": 0,
                "duration_ms": 0,
                "failure_reason": None,
                "selectors": [],
            }
            for key in all_keys
        ]
        _update(
            job["job_id"],
            codegen_enhance_steps=enhance_steps,
            current_step="AI 增强项目代码",
        )
    else:
        _update(job["job_id"], codegen_enhance_steps=[])

    def _event_key(event: Dict[str, str]) -> Optional[str]:
        if "step" in event and event["step"]:
            return str(event["step"])
        if "file" in event:
            f = event["file"]
            if f and "::" in f:
                return f.split("::", 1)[1]
            if f == "README.md":
                return README_KEY
        return None

    def update_enhance_step(event: Dict[str, str]) -> None:
        if not enhance_steps:
            return
        key = _event_key(event)
        if not key:
            return
        for item in enhance_steps:
            if item["key"] == key:
                item["status"] = event.get("status", item["status"])
                for optional in (
                    "summary",
                    "attempts",
                    "duration_ms",
                    "failure_reason",
                    "selectors",
                ):
                    if optional in event and event[optional] not in (None, ""):
                        item[optional] = event[optional]
                break
        current_name = next(
            (item["name"] for item in enhance_steps if item["key"] == key),
            "项目代码",
        )
        _update(
            job["job_id"],
            codegen_enhance_steps=enhance_steps,
            current_step=f"AI 增强项目代码：{current_name}",
        )

    result = enhance_project(
        job,
        planning,
        job_dir / "generated_project",
        job_dir / ".enhancer_backup",
        progress_callback=update_enhance_step,
    )
    enhancement_payload = result.model_dump()
    _json_write(job_dir / "enhancement.json", enhancement_payload)
    # 同步落盘 ui_enhancement.json：保留可独立审计的 UI 步骤摘要
    if result.ui_steps or result.ui_plan is not None:
        ui_enhancement = {
            "schema_version": "1.0",
            "generated_at": _now(),
            "mode": {
                "requested": result.requested_mode,
                "actual": result.actual_mode,
            },
            "model": result.model,
            "ui_plan": result.ui_plan,
            "ui_steps": result.ui_steps,
            "fallback_reason": result.fallback_reason,
            "summary": result.summary,
        }
        _json_write(job_dir / "ui_enhancement.json", ui_enhancement)
    _update(
        job["job_id"],
        codegen_actual_mode=result.actual_mode,
        codegen_model=result.model,
        codegen_summary=result.summary,
        codegen_changed_files=result.changed_files,
        codegen_fallback_reason=result.fallback_reason,
    )


def validate_generated_project(job: Dict[str, Any], job_dir: Path) -> None:
    try:
        run_generated_project(job_dir)
        validation = _json_read(job_dir / "run_validation.json")
        _update(
            job["job_id"],
            run_status=(
                "verified"
                if validation["maven_test"] == "passed"
                else "structure_verified"
            ),
            run_validation=validation,
            demo_url="http://127.0.0.1:9002",
            swagger_url="http://127.0.0.1:9001/swagger-ui/index.html",
        )
    except Exception as exc:
        status = get_job(job["job_id"]) or {}
        if (
            job.get("codegen_mode") == "auto"
            and status.get("codegen_actual_mode") == "llm"
        ):
            restored = restore_enhancement(
                job_dir / "generated_project",
                job_dir / ".enhancer_backup",
            )
            reason = f"{type(exc).__name__}: {str(exc)[:600]}"
            _update(
                job["job_id"],
                codegen_actual_mode="template",
                codegen_fallback_reason=(
                    f"增强代码验证失败，已回滚 {len(restored)} 个文件: "
                    f"{reason}"
                ),
            )
            enhancement = _json_read(job_dir / "enhancement.json")
            enhancement["actual_mode"] = "template"
            enhancement["fallback_reason"] = (
                f"验证失败并回滚: {reason}"
            )
            _json_write(job_dir / "enhancement.json", enhancement)
            run_generated_project(job_dir)
            validation = _json_read(job_dir / "run_validation.json")
            _update(
                job["job_id"],
                run_status=(
                    "verified"
                    if validation["maven_test"] == "passed"
                    else "structure_verified"
                ),
                run_validation=validation,
                demo_url="http://127.0.0.1:9002",
                swagger_url="http://127.0.0.1:9001/swagger-ui/index.html",
            )
            return
        raise


def _write(path: Path, content: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(content.strip() + "\n", encoding="utf-8")


def generate_project(job_dir: Path) -> None:
    generate_java_project(job_dir)
    return
    planning = _json_read(job_dir / "planning.json")
    root = job_dir / "generated_project"
    frontend = root / "frontend"
    backend = root / "backend"
    name = planning["software_name"]
    modules = planning["modules"]
    _write(
        frontend / "package.json",
        """{
  "name": "copyright-demo-frontend",
  "private": true,
  "version": "0.1.0",
  "type": "module",
  "scripts": {"dev": "vite --host 127.0.0.1", "build": "vite build"},
  "dependencies": {"@vitejs/plugin-vue": "^5.2.1", "vite": "^6.0.5", "vue": "^3.5.13"},
  "devDependencies": {}
}""",
    )
    _write(
        frontend / "vite.config.js",
        """import { defineConfig } from 'vite'
import vue from '@vitejs/plugin-vue'
import { dirname } from 'node:path'
import { fileURLToPath } from 'node:url'

export default defineConfig({
  root: dirname(fileURLToPath(import.meta.url)),
  plugins: [vue()]
})
""",
    )
    _write(
        frontend / "index.html",
        '<!doctype html><html lang="zh-CN"><head><meta charset="UTF-8"><meta name="viewport" content="width=device-width,initial-scale=1.0"><title>'
        + name
        + '</title></head><body><div id="app"></div><script type="module" src="/src/main.js"></script></body></html>',
    )
    _write(
        frontend / "src/main.js",
        """import { createApp } from 'vue'
import './style.css'
import App from './App.vue'
createApp(App).mount('#app')
""",
    )
    module_json = json.dumps(modules, ensure_ascii=False)
    mock_rows = {
        module["key"]: [
            [f"{field}示例一" for field in module["fields"]],
            [f"{field}示例二" for field in module["fields"]],
        ]
        for module in modules
    }
    mock_rows_json = json.dumps(mock_rows, ensure_ascii=False)
    app_vue = """<script setup>
import { computed, ref } from 'vue'
const softwareName = %s
const modules = %s
const active = ref('dashboard')
const loggedIn = ref(false)
const showModal = ref(false)
const rows = %s
const current = computed(() => modules.find(m => m.key === active.value))
</script>
<template>
  <div v-if="!loggedIn" class="login-page">
    <div class="login-brand"><h1>{{softwareName}}</h1><p>业务运营管理平台</p></div>
    <div class="login-card"><h2>欢迎登录</h2><p>请输入管理员账号进入系统</p><label>用户名</label><input value="admin"><label>密码</label><input type="password" value="123456"><button @click="loggedIn=true">登录系统</button><small>演示账号：admin / 123456</small></div>
  </div>
  <div v-else class="shell">
    <aside><h2>{{ softwareName }}</h2><div class="brand-sub">智慧运营中心</div>
      <button :class="{on:active==='dashboard'}" @click="active='dashboard'">运营首页</button>
      <button v-for="m in modules" :key="m.key" :data-module-key="m.key" :class="{on:active===m.key}" @click="active=m.key">{{m.name}}</button>
    </aside>
    <main>
      <header><div><b>{{ active==='dashboard'?'运营首页':current.name }}</b><span>数据更新时间：2026-06-10 10:30</span></div><div class="user">管理员</div></header>
      <section v-if="active==='dashboard'">
        <div class="hero"><div><p>系统运行概览</p><h1>{{softwareName}}运行正常</h1><small>当前已启用 {{modules.length}} 个业务模块</small></div><div class="rate">{{modules.length}}<small>业务模块</small></div></div>
        <div class="cards"><article v-for="(module,index) in modules.slice(0,4)"><span>{{module.name}}</span><b>{{128 + index * 37}}</b><i>数据状态正常</i></article></div>
        <div class="panel"><h3>近七日业务趋势</h3><div class="chart"><i v-for="h in [48,60,52,72,66,84,76]" :style="{height:h+'%%'}"></i></div></div>
      </section>
      <section v-else>
        <div class="toolbar"><div><h2>{{current.name}}</h2><p>{{current.description || ('统一维护'+current.name+'相关业务数据')}}</p></div><button class="primary" @click="showModal=true">+ 新增记录</button></div>
        <div class="filters"><input placeholder="请输入关键字搜索"><button>查询</button><button>重置</button></div>
        <div class="panel table"><table><thead><tr><th>序号</th><th v-for="f in current.fields">{{f}}</th><th>操作</th></tr></thead><tbody><tr v-for="(row,i) in rows[active]"><td>{{i+1}}</td><td v-for="cell in row">{{cell}}</td><td><a>编辑</a><a>详情</a></td></tr></tbody></table></div>
      </section>
    </main>
    <div v-if="showModal" class="modal-mask"><div class="modal"><h2>新增{{current.name}}记录</h2><p>请填写以下业务信息</p><label v-for="field in current.fields">{{field}}<input :placeholder="'请输入'+field"></label><div><button @click="showModal=false">取消</button><button class="primary" @click="showModal=false">保存</button></div></div></div>
  </div>
</template>
""" % (json.dumps(name, ensure_ascii=False), module_json, mock_rows_json)
    _write(frontend / "src/App.vue", app_vue)
    _write(
        frontend / "src/style.css",
        """*{box-sizing:border-box}body{margin:0;font-family:"Microsoft YaHei",Arial;color:#1f2d3d;background:#f3f6fb}.login-page{min-height:100vh;background:linear-gradient(120deg,#0d3268,#2186c9);display:flex;align-items:center;justify-content:center;gap:130px;color:white}.login-brand h1{font-size:38px;margin:0 0 10px}.login-brand p{opacity:.7;letter-spacing:4px}.login-card{width:380px;background:white;color:#24354b;padding:38px;border-radius:14px;box-shadow:0 22px 60px #06224655}.login-card h2{margin:0}.login-card p,.login-card small{color:#8b98a9}.login-card label,.modal label{display:block;font-size:13px;margin:18px 0 7px}.login-card input,.modal input{width:100%;padding:11px;border:1px solid #d6dfeb;border-radius:6px}.login-card button{width:100%;padding:12px;margin:22px 0 14px;border:0;border-radius:6px;background:#176bc1;color:white}.shell{display:flex;min-height:100vh}aside{width:230px;background:linear-gradient(180deg,#123b73,#0a2750);color:white;padding:25px 16px}aside h2{font-size:20px;margin:0 8px 4px}.brand-sub{font-size:12px;opacity:.6;margin:0 8px 28px}aside button{display:block;width:100%;border:0;background:transparent;color:#cbd8ea;text-align:left;padding:13px 16px;margin:4px 0;border-radius:7px;font-size:14px;cursor:pointer}aside button.on,aside button:hover{background:#1d579c;color:white}main{flex:1}header{height:68px;background:white;display:flex;justify-content:space-between;align-items:center;padding:0 30px;border-bottom:1px solid #e7ecf3}header b{font-size:18px}header span{font-size:12px;color:#8795a8;margin-left:20px}.user{background:#edf4ff;color:#2868b2;padding:8px 15px;border-radius:18px}section{padding:24px 30px}.hero{background:linear-gradient(120deg,#1769c2,#25a2d8);color:white;border-radius:12px;padding:28px 35px;display:flex;justify-content:space-between;align-items:center;box-shadow:0 10px 24px #1d72b533}.hero p{margin:0}.hero h1{margin:8px 0;font-size:27px}.hero small{opacity:.8}.rate{font-size:36px;font-weight:bold;text-align:center}.rate small{display:block;font-size:13px;font-weight:normal}.cards{display:grid;grid-template-columns:repeat(4,1fr);gap:18px;margin:20px 0}.cards article,.panel{background:white;border-radius:10px;padding:22px;box-shadow:0 3px 12px #193b6010}.cards span{display:block;color:#7f8da0}.cards b{display:block;font-size:28px;margin:12px 0}.cards i{font-style:normal;color:#3d9b72;font-size:12px}.panel h3{margin:0}.chart{height:210px;display:flex;align-items:flex-end;gap:34px;padding:30px 45px 0;border-bottom:1px solid #dfe6ef}.chart i{flex:1;background:linear-gradient(#51a6e8,#1970c4);border-radius:5px 5px 0 0}.toolbar{display:flex;justify-content:space-between;align-items:center}.toolbar h2{margin:0}.toolbar p{color:#8795a8}.primary{background:#176bc1!important;color:white;border:0!important}.toolbar button,.filters button{padding:10px 18px;border:1px solid #d7e0eb;border-radius:5px;background:white}.filters{background:white;padding:18px;margin:16px 0;border-radius:8px}.filters input{width:280px;padding:10px;border:1px solid #d7e0eb;border-radius:5px;margin-right:10px}.table{padding:0;overflow:hidden}table{width:100%;border-collapse:collapse}th,td{text-align:left;padding:16px;border-bottom:1px solid #edf0f4}th{background:#f8fafc;color:#657489}td a{color:#176bc1;margin-right:14px;font-size:13px}.modal-mask{position:fixed;inset:0;background:#10243d88;display:grid;place-items:center}.modal{width:480px;background:white;border-radius:12px;padding:28px;box-shadow:0 20px 60px #10243d55}.modal h2{margin:0}.modal p{color:#8290a2}.modal>div{text-align:right;margin-top:22px}.modal>div button{padding:10px 22px;margin-left:10px;border:1px solid #d8e0ea;background:white;border-radius:6px}""",
    )
    _write(
        backend / "requirements.txt",
        "fastapi==0.115.6\nuvicorn[standard]==0.32.1",
    )
    route_definitions = []
    schema_statements = []
    database_tables = planning["database_tables"]
    for index, module in enumerate(modules):
        key = module["key"]
        table_name = database_tables[index % len(database_tables)]
        route_definitions.append(
            f'''
@app.get("/api/{key}")
def list_{key}():
    return rows("{table_name}")

@app.post("/api/{key}")
def create_{key}(payload: dict):
    return {{"id": 3, **payload}}
'''
        )
    for index, table_name in enumerate(database_tables):
        module = modules[index % len(modules)]
        columns = ", ".join(
            f"field_{field_index} TEXT"
            for field_index, _ in enumerate(module["fields"], start=1)
        )
        schema_statements.append(
            f"CREATE TABLE IF NOT EXISTS {table_name} (id INTEGER PRIMARY KEY, {columns});"
        )
        for row_index in range(1, 3):
            values = [
                f"{field}示例{'一' if row_index == 1 else '二'}"
                for field in module["fields"]
            ]
            sql_values = ", ".join(
                "'" + value.replace("'", "''") + "'" for value in values
            )
            schema_statements.append(
                f"INSERT INTO {table_name} VALUES ({row_index}, {sql_values});"
            )
    routes_code = "\n".join(route_definitions)
    schema_sql = "\n".join(schema_statements)
    _write(
        backend / "main.py",
        """import sqlite3
from pathlib import Path
from fastapi import FastAPI
from fastapi.middleware.cors import CORSMiddleware

app = FastAPI(title=%s)
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_methods=["*"], allow_headers=["*"])
DB = Path(__file__).with_name("app.db")

def rows(table):
    with sqlite3.connect(DB) as conn:
        conn.row_factory = sqlite3.Row
        return [dict(row) for row in conn.execute("select * from " + table)]

@app.get("/api/health")
def health(): return {"status": "ok"}

@app.get("/api/dashboard")
def dashboard():
    return {"module_count": %s, "status": "running", "today_records": 328, "completion_rate": 78.6}

%s
""" % (json.dumps(name, ensure_ascii=False), len(modules), routes_code),
    )
    _write(backend / "schema.sql", schema_sql)
    with __import__("sqlite3").connect(backend / "app.db") as conn:
        conn.executescript((backend / "schema.sql").read_text(encoding="utf-8"))
    _write(
        root / "README.md",
        f"""# {name}

本项目由 AI软著工厂 Phase 3 生成，基础结构来自固定模板，并可经过受约束的 AI 代码增强。

## 后端
```powershell
cd "{(root / 'backend').resolve()}"
python -m pip install -r requirements.txt
python -m uvicorn main:app --port 9001
```

## 前端
```powershell
cd "{(root / 'frontend').resolve()}"
npm.cmd install
npm.cmd run dev -- --port 9002
```
""",
    )


def _wait_port(port: int, timeout: int = 30) -> None:
    started = time.time()
    while time.time() - started < timeout:
        with socket.socket() as sock:
            if sock.connect_ex(("127.0.0.1", port)) == 0:
                return
        time.sleep(0.4)
    raise RuntimeError(f"端口 {port} 启动超时")


def _free_port() -> int:
    with socket.socket() as sock:
        sock.bind(("127.0.0.1", 0))
        return int(sock.getsockname()[1])


def _terminate_pid(pid: Optional[int]) -> None:
    if not pid:
        return
    try:
        if os.name == "nt":
            subprocess.run(
                ["taskkill", "/PID", str(pid), "/T", "/F"],
                capture_output=True,
                timeout=15,
            )
        else:
            os.kill(pid, 15)
    except (OSError, subprocess.SubprocessError):
        pass


def _port_open(port: Optional[int]) -> bool:
    if not port:
        return False
    with socket.socket() as sock:
        sock.settimeout(0.5)
        return sock.connect_ex(("127.0.0.1", int(port))) == 0


def _java_sources_mtime(backend: Path) -> float:
    sources = list((backend / "src").rglob("*.java")) if (backend / "src").exists() else []
    pom = backend / "pom.xml"
    latest = pom.stat().st_mtime if pom.exists() else 0
    for path in sources:
        latest = max(latest, path.stat().st_mtime)
    return latest


def _find_existing_jar(backend: Path) -> Optional[Path]:
    target = backend / "target"
    if not target.exists():
        return None
    jars = sorted(
        (p for p in target.glob("*.jar") if not p.name.endswith(".original")),
        key=lambda p: p.stat().st_mtime,
    )
    return jars[-1] if jars else None


def _launch_demo(
    job_id: str,
    job_dir: Path,
    backend_port: int,
    frontend_port: int,
) -> Tuple[subprocess.Popen, subprocess.Popen, str]:
    project = job_dir / "generated_project"
    backend = project / "backend"
    frontend = project / "frontend"
    maven = _maven_command()
    if not maven:
        raise RuntimeError("未找到 Maven，无法启动 Spring Boot Demo")
    maven_version = _maven_version(maven)

    logs = job_dir / "logs"
    logs.mkdir(exist_ok=True)
    build_log_path = logs / "demo_build.log"

    # 复用已经构建的 JAR：源代码/POM 没更新就不重跑 mvn
    existing_jar = _find_existing_jar(backend)
    src_mtime = _java_sources_mtime(backend)
    if existing_jar and existing_jar.stat().st_mtime >= src_mtime:
        _set_demo_stage(
            job_id,
            stage="building",
            stage_detail="复用已构建的 JAR，跳过 mvn package",
            maven_version=maven_version,
            jar_reused=True,
        )
        jars = [existing_jar]
    else:
        _set_demo_stage(
            job_id,
            stage="building",
            stage_detail="正在执行 mvn package -DskipTests",
            maven_version=maven_version,
            jar_reused=False,
        )
        with open(build_log_path, "a", encoding="utf-8") as build_log:
            build_log.write(f"\n===== mvn package @ {_now()} =====\n")
            mvn_env = _maven_subprocess_env()
            package_result = subprocess.run(
                [maven, "package", "-DskipTests"],
                cwd=str(backend),
                capture_output=True,
                text=True,
                encoding="utf-8",
                errors="replace",
                timeout=300,
                env=mvn_env,
            )
            build_log.write(package_result.stdout or "")
            if package_result.stderr:
                build_log.write("\n--- stderr ---\n")
                build_log.write(package_result.stderr)
            build_log.write(f"\nexit={package_result.returncode}\n")
        if package_result.returncode != 0:
            detail = (package_result.stderr or package_result.stdout)[-1500:]
            raise RuntimeError("Spring Boot JAR 构建失败: " + detail)
        jars = sorted(
            path
            for path in (backend / "target").glob("*.jar")
            if not path.name.endswith(".original")
        )
        if not jars:
            raise RuntimeError("Maven 构建完成但未找到可执行 JAR")

    _set_demo_stage(
        job_id,
        stage="starting",
        stage_detail="正在启动 Spring Boot 与 Vite",
    )
    backend_output = open(logs / "backend.log", "a", encoding="utf-8")
    frontend_output = open(logs / "frontend.log", "a", encoding="utf-8")
    creationflags = subprocess.CREATE_NO_WINDOW if os.name == "nt" else 0
    backend_process = subprocess.Popen(
        [
            shutil.which("java") or "java",
            "-jar",
            str(jars[-1]),
            "--spring.profiles.active=demo",
            f"--server.port={backend_port}",
        ],
        cwd=str(backend),
        stdout=backend_output,
        stderr=subprocess.STDOUT,
        creationflags=creationflags,
    )
    frontend_env = os.environ.copy()
    frontend_env["VITE_BACKEND_TARGET"] = f"http://127.0.0.1:{backend_port}"
    frontend_process = subprocess.Popen(
        [_npm_command(), "run", "dev", "--", "--port", str(frontend_port)],
        cwd=str(frontend),
        stdout=frontend_output,
        stderr=subprocess.STDOUT,
        env=frontend_env,
        creationflags=creationflags,
    )
    backend_output.close()
    frontend_output.close()
    try:
        _wait_port(backend_port, 90)
        _wait_port(frontend_port, 40)
    except Exception:
        _terminate_pid(backend_process.pid)
        _terminate_pid(frontend_process.pid)
        raise
    return backend_process, frontend_process, maven_version


def stop_online_demo(job_id: str) -> Dict[str, Any]:
    runtime_path = OUTPUT_ROOT / job_id / "demo_runtime.json"
    if not runtime_path.exists():
        return {"status": "stopped", "stage": "stopped"}
    runtime = _json_read(runtime_path)
    if _port_open(runtime.get("backend_port")):
        _terminate_pid(runtime.get("backend_pid"))
    if _port_open(runtime.get("frontend_port")):
        _terminate_pid(runtime.get("frontend_pid"))
    runtime["status"] = "stopped"
    runtime["stage"] = "stopped"
    runtime["stopped_at"] = _now()
    _json_write(runtime_path, runtime)
    _update(job_id, run_status="stopped")
    return runtime


def start_online_demo(job: Dict[str, Any], job_dir: Path) -> Dict[str, Any]:
    stop_online_demo(job["job_id"])
    # 立刻写入 starting 状态，让前端轮询时能区分 "正在启动"
    # 与"已停止"，并展示当前 stage。
    initial_runtime = {
        "status": "starting",
        "stage": "queued",
        "stage_detail": "排队中…",
        "started_at": _now(),
        "expires_at": (datetime.now() + timedelta(hours=1)).isoformat(
            timespec="seconds"
        ),
    }
    _json_write(job_dir / "demo_runtime.json", initial_runtime)
    _update(job["job_id"], run_status="starting")

    backend_port = _free_port()
    frontend_port = _free_port()
    backend_process, frontend_process, maven_version = _launch_demo(
        job["job_id"],
        job_dir,
        backend_port,
        frontend_port,
    )
    runtime = {
        "status": "running",
        "stage": "running",
        "stage_detail": "Demo 已就绪",
        "maven_version": maven_version,
        "backend_pid": backend_process.pid,
        "frontend_pid": frontend_process.pid,
        "backend_port": backend_port,
        "frontend_port": frontend_port,
        "demo_url": f"http://127.0.0.1:{frontend_port}",
        "swagger_url": f"http://127.0.0.1:{backend_port}/swagger-ui/index.html",
        "started_at": _now(),
        "expires_at": (datetime.now() + timedelta(hours=1)).isoformat(timespec="seconds"),
    }
    _json_write(job_dir / "demo_runtime.json", runtime)
    _update(
        job["job_id"],
        run_status="running",
        demo_url=runtime["demo_url"],
        swagger_url=runtime["swagger_url"],
    )
    return runtime


def demo_runtime(job_id: str) -> Dict[str, Any]:
    path = OUTPUT_ROOT / job_id / "demo_runtime.json"
    if not path.exists():
        return {"status": "stopped"}
    runtime = _json_read(path)
    expires_at = runtime.get("expires_at")
    if runtime.get("status") == "running" and expires_at:
        if datetime.now() >= datetime.fromisoformat(expires_at):
            return stop_online_demo(job_id)
        if not (
            _port_open(runtime.get("backend_port"))
            and _port_open(runtime.get("frontend_port"))
        ):
            runtime["status"] = "stopped"
            runtime["stopped_at"] = _now()
            runtime["stop_reason"] = "Demo 进程已退出"
            _json_write(path, runtime)
            _update(job_id, run_status="stopped")
    return runtime


def _npm_command() -> str:
    return "npm.cmd" if os.name == "nt" else "npm"


def _maven_command() -> Optional[str]:
    """选择 Maven 可执行文件。

    优先使用 IntelliJ 自带 mvn.cmd（与生成项目兼容性好、避免 PluginContainerException），
    其次才回退到系统 PATH 中的 mvn。
    """
    if os.name == "nt":
        candidates = [
            Path(r"C:\Program Files\JetBrains\IntelliJ IDEA 2025.1\plugins\maven\lib\maven3\bin\mvn.cmd"),
            Path(r"C:\Program Files\JetBrains\IntelliJ IDEA 2024.1\plugins\maven\lib\maven3\bin\mvn.cmd"),
            Path(r"C:\Program Files\JetBrains\IntelliJ IDEA 2023.1\plugins\maven\lib\maven3\bin\mvn.cmd"),
        ]
        found = next((path for path in candidates if path.exists()), None)
        if found:
            return str(found)
    command = shutil.which("mvn.cmd" if os.name == "nt" else "mvn")
    return command


def _maven_subprocess_env() -> Dict[str, str]:
    """构造调用 mvn 时的环境变量。

    IntelliJ 自带 mvn.cmd 内部默认走 Java 8 启动器，但生成项目依赖 mybatis-spring-3.0.3+
    已经是 Java 17 字节码（version 61.0）。Java 8 无法读取 61.0 类文件，编译会报
    "类文件具有的版本 61.0, 应为 52.0"。

    解决：自动探测本机 JDK 17（系统安装/IntelliJ 自带 JBR），并把 JAVA_HOME 指向它，
    这样 mvn 会用 Java 17 跑 javac 编译，避免版本不匹配。
    """
    env = os.environ.copy()
    if os.name != "nt":
        return env
    candidates = [
        Path(r"D:\Program Files\Java\jdk-17"),
        Path(r"C:\Program Files\Java\jdk-17"),
        Path(r"D:\Program Files\JetBrains\IntelliJ IDEA 2023.1\jbr"),
        Path(r"C:\Program Files\JetBrains\IntelliJ IDEA 2023.1\jbr"),
    ]
    java17 = next((p for p in candidates if (p / "bin" / "java.exe").exists()), None)
    if java17:
        env["JAVA_HOME"] = str(java17)
        env["PATH"] = f"{java17}\\bin;{env.get('PATH', '')}"
    return env


def _maven_version(maven: str) -> str:
    try:
        result = subprocess.run(
            [maven, "-v"],
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            timeout=15,
        )
    except (OSError, subprocess.SubprocessError, subprocess.TimeoutExpired) as exc:
        return f"{maven} (version unknown: {type(exc).__name__})"
    text = (result.stdout or result.stderr or "").strip()
    return text.splitlines()[0] if text else maven


def run_generated_project(job_dir: Path) -> None:
    project = job_dir / "generated_project"
    frontend = project / "frontend"
    backend = project / "backend"
    validation = {
        "frontend_build": "pending",
        "backend_structure": "pending",
        "maven_test": "unavailable",
        "maven_reason": None,
    }
    if not (frontend / "node_modules").exists():
        install_timeout = int(os.getenv("NPM_INSTALL_TIMEOUT", "600"))
        install = subprocess.run(
            [_npm_command(), "install", "--no-audit", "--no-fund", "--prefer-offline"],
            cwd=str(frontend),
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            timeout=install_timeout,
        )
        if install.returncode != 0:
            raise RuntimeError("前端依赖安装失败: " + install.stderr[-500:])
    build = subprocess.run(
        [_npm_command(), "run", "build"],
        cwd=str(frontend),
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
        timeout=120,
    )
    if build.returncode != 0:
        detail = (build.stderr or build.stdout)[-1000:]
        raise RuntimeError("前端构建失败: " + detail)
    validation["frontend_build"] = "passed"

    required = [
        backend / "pom.xml",
        backend / "src/main/resources/application.yml",
        project / "sql/init.sql",
    ]
    java_files = list((backend / "src/main/java").rglob("*.java"))
    if any(not path.exists() for path in required) or not java_files:
        raise RuntimeError("Java 项目结构不完整")
    validation["backend_structure"] = "passed"

    maven = _maven_command()
    wrapper = backend / ("mvnw.cmd" if os.name == "nt" else "mvnw")
    command = [maven, "test"] if maven else ([str(wrapper), "test"] if wrapper.exists() else None)
    if command:
        result = subprocess.run(
            command,
            cwd=str(backend),
            env=_maven_subprocess_env(),
            capture_output=True,
            text=True,
            encoding="utf-8",
            errors="replace",
            timeout=300,
        )
        if result.returncode != 0:
            detail = (result.stderr or result.stdout)[-1500:]
            raise RuntimeError("Maven 测试失败: " + detail)
        validation["maven_test"] = "passed"
    else:
        validation["maven_reason"] = "当前主机未安装 Maven，且生成项目中没有可用 Maven Wrapper"
    _json_write(job_dir / "run_validation.json", validation)


def _edge_path() -> Optional[str]:
    candidates = [
        Path(r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe"),
        Path(r"C:\Program Files\Microsoft\Edge\Application\msedge.exe"),
    ]
    return str(next((path for path in candidates if path.exists()), "")) or None


def _safe_filename(value: str) -> str:
    return re.sub(r'[<>:"/\\|?*]+', "-", value).strip(" .") or "page"


def _wait_for_settle(page: Any, kind: str) -> Dict[str, Any]:
    """按截图类型等 Element Plus 动画 / 异步加载完成。

    返回 ``{"strategy": str, "duration_ms": int, "retried": bool}`` 供 manifest 记录。
    """
    import time as _time

    start = _time.monotonic()

    def _elapsed() -> int:
        return int((_time.monotonic() - start) * 1000)

    if kind == "login":
        # 登录页通常只等 networkidle；保留最低 200ms 让登录卡淡入完成
        page.wait_for_load_state("networkidle")
        page.wait_for_timeout(200)
        return {"strategy": "login_idle", "duration_ms": _elapsed(), "retried": False}

    if kind == "dashboard":
        # Dashboard 跳转后等 networkidle + KPI 容器可见
        page.wait_for_load_state("networkidle")
        try:
            page.wait_for_selector(".kpi-card, .kpi-grid, .hero", timeout=4000)
        except Exception:
            pass
        return {"strategy": "dashboard_idle", "duration_ms": _elapsed(), "retried": False}

    if kind == "list":
        # 列表页：等网络空闲 + 首行渲染
        page.wait_for_load_state("networkidle")
        try:
            page.locator(".el-table__row, .el-empty, .module-page .empty").first.wait_for(
                timeout=4000
            )
        except Exception:
            pass
        # 额外 250ms 让 loading 蒙层完全褪去
        page.wait_for_timeout(250)
        return {"strategy": "list_idle+row", "duration_ms": _elapsed(), "retried": False}

    if kind == "dialog":
        # 对话框：等 el-dialog 出现 + overlay opacity >= 0.95 + wrapper transform 归位
        try:
            page.wait_for_selector(".el-dialog", timeout=5000)
        except Exception:
            return {"strategy": "dialog_timeout", "duration_ms": _elapsed(), "retried": False}
        attempts = 0
        for _ in range(2):  # 一次重试机会
            attempts += 1
            try:
                page.wait_for_function(
                    """() => {
                        const overlay = document.querySelector('.el-overlay');
                        const wrapper = document.querySelector('.el-dialog__wrapper');
                        if (!overlay || !wrapper) return false;
                        const s = getComputedStyle(overlay);
                        const w = getComputedStyle(wrapper);
                        const opacity = parseFloat(s.opacity || '0');
                        const transform = w.transform || '';
                        const settled = opacity >= 0.95
                            && (transform === 'none'
                                || transform.includes('matrix(1, 0, 0, 1, 0, 0)'));
                        return settled;
                    }""",
                    timeout=3000,
                )
                break
            except Exception:
                # 兜底：固定 400ms 等待
                page.wait_for_timeout(400)
        # 再加 150ms 让内部表单 v-model 同步完成
        page.wait_for_timeout(150)
        return {
            "strategy": "dialog_anim",
            "duration_ms": _elapsed(),
            "retried": attempts > 1,
        }

    if kind == "dialog_close":
        # 关闭对话框：等 el-dialog 完全从 DOM 移除
        try:
            page.wait_for_selector(
                ".el-dialog", state="detached", timeout=4000
            )
        except Exception:
            page.wait_for_timeout(400)
        page.wait_for_timeout(150)
        return {"strategy": "dialog_detached", "duration_ms": _elapsed(), "retried": False}

    return {"strategy": "none", "duration_ms": _elapsed(), "retried": False}


def _stretch_overlay_to_page(page: Any) -> bool:
    """把 .el-overlay 从 position:fixed 改为 absolute 并撑满整页高度。

    ``full_page=True`` 截图时，``position: fixed`` 元素只会渲染在视口顶部
    那一段（~900px），视口下方的页面内容会原样进入图片，导致出现
    "主页面表格和子菜单（对话框）上下叠在一起"的假象。

    修复策略：截图前用 JS 把 overlay 改成 ``position: absolute; top: 0;``
    高度等于 ``document.documentElement.scrollHeight``，让蒙层覆盖整页。
    对话框本身仍由 Element Plus 居中定位，不动。

    Returns ``True`` 当 overlay 被改写过，``False`` 表示 DOM 里找不到蒙层。
    """
    return bool(
        page.evaluate(
            """() => {
                const overlay = document.querySelector('.el-overlay');
                if (!overlay) return false;
                const pageHeight = document.documentElement.scrollHeight;
                overlay.style.position = 'absolute';
                overlay.style.top = '0';
                overlay.style.left = '0';
                overlay.style.right = '0';
                overlay.style.width = '100%';
                overlay.style.height = pageHeight + 'px';
                overlay.style.minHeight = '100vh';
                return true;
            }"""
        )
    )


def capture_screenshots(job_dir: Path) -> None:
    from playwright.sync_api import sync_playwright

    planning = _json_read(job_dir / "planning.json")
    screenshot_modules = planning["modules"]
    screenshot_dir = job_dir / "screenshots"
    screenshot_dir.mkdir(exist_ok=True)
    manifest: List[Dict[str, Any]] = []
    runtime_path = job_dir / "demo_runtime.json"
    runtime = _json_read(runtime_path) if runtime_path.exists() else {}
    borrowed_demo = (
        runtime.get("status") == "running"
        and _port_open(runtime.get("backend_port"))
        and _port_open(runtime.get("frontend_port"))
    )
    backend_process = None
    frontend_process = None
    if borrowed_demo:
        frontend_port = int(runtime["frontend_port"])
    else:
        backend_port = _free_port()
        frontend_port = _free_port()
        job_id = job_dir.name
        backend_process, frontend_process, _ = _launch_demo(
            job_id,
            job_dir,
            backend_port,
            frontend_port,
        )
    try:
        with sync_playwright() as playwright:
            browser = playwright.chromium.launch(headless=True, executable_path=_edge_path())
            page = browser.new_page(viewport={"width": 1440, "height": 900}, device_scale_factor=1)
            page.goto(f"http://127.0.0.1:{frontend_port}", wait_until="networkidle")
            login_path = screenshot_dir / "01-login.png"
            page.screenshot(path=str(login_path), full_page=True)
            login_settle = _wait_for_settle(page, "login")
            manifest.append({
                "kind": "login",
                "file": login_path.name,
                "label": "系统登录",
                "wait_strategy": login_settle["strategy"],
                "duration_ms": login_settle["duration_ms"],
            })
            page.locator(".login-card button").click()
            dashboard_path = screenshot_dir / "02-dashboard.png"
            page.screenshot(path=str(dashboard_path), full_page=True)
            dashboard_settle = _wait_for_settle(page, "dashboard")
            manifest.append({
                "kind": "dashboard",
                "file": dashboard_path.name,
                "label": "系统首页",
                "wait_strategy": dashboard_settle["strategy"],
                "duration_ms": dashboard_settle["duration_ms"],
            })
            index = 3
            for module in screenshot_modules:
                label = module["name"]
                page.locator(f'[data-module-key="{module["key"]}"]').click()
                page.locator(f'section[data-module-key="{module["key"]}"]').wait_for()
                safe_label = _safe_filename(label)
                list_path = screenshot_dir / f"{index:02d}-{safe_label}-list.png"
                page.screenshot(
                    path=str(list_path),
                    full_page=True,
                )
                list_settle = _wait_for_settle(page, "list")
                manifest.append(
                    {
                        "kind": "module_list",
                        "file": list_path.name,
                        "module_key": module["key"],
                        "module_name": label,
                        "label": f"{label}功能页",
                        "wait_strategy": list_settle["strategy"],
                        "duration_ms": list_settle["duration_ms"],
                    }
                )
                index += 1
                page.locator('[data-action="create"]').click()
                dialog_settle = _wait_for_settle(page, "dialog")
                # full_page=True 时 fixed overlay 只覆盖顶部视口，
                # 会让视口下方的页面表格"裸露"在截图里。
                # 截图前把 overlay 改成 absolute + 整页高度，让蒙层真的把全页盖住。
                overlay_stretched = _stretch_overlay_to_page(page)
                form_path = screenshot_dir / f"{index:02d}-{safe_label}-create.png"
                page.screenshot(path=str(form_path), full_page=True)
                manifest.append(
                    {
                        "kind": "module_create",
                        "file": form_path.name,
                        "module_key": module["key"],
                        "module_name": label,
                        "label": f"{label}新增表单",
                        "wait_strategy": dialog_settle["strategy"],
                        "duration_ms": dialog_settle["duration_ms"],
                        "retried_after_fix": dialog_settle["retried"],
                        "overlay_full_page": overlay_stretched,
                    }
                )
                page.locator(".el-dialog__footer button").first.click()
                _wait_for_settle(page, "dialog_close")
                index += 1
            browser.close()
        _json_write(job_dir / "screenshot_manifest.json", {"screenshots": manifest})
    finally:
        if not borrowed_demo:
            _terminate_pid(backend_process.pid if backend_process else None)
            _terminate_pid(frontend_process.pid if frontend_process else None)


CODE_EXTENSIONS = {".java", ".xml", ".yml", ".yaml", ".vue", ".js", ".ts", ".html", ".css", ".sql"}
EXCLUDED = {"node_modules", "dist", "venv", ".git", "__pycache__"}


def _source_files(root: Path) -> Iterable[Path]:
    for path in sorted(root.rglob("*")):
        if path.is_file() and path.suffix.lower() in CODE_EXTENSIONS:
            if not any(part in EXCLUDED for part in path.parts):
                yield path


def analyze_code(job_dir: Path) -> None:
    root = job_dir / "generated_project"
    stats = {"total_lines": 0, "frontend_lines": 0, "backend_lines": 0, "sql_lines": 0, "files": []}
    for path in _source_files(root):
        lines = len(path.read_text(encoding="utf-8", errors="ignore").splitlines())
        relative = path.relative_to(root).as_posix()
        stats["total_lines"] += lines
        if relative.startswith("frontend/"):
            stats["frontend_lines"] += lines
        if relative.startswith("backend/"):
            stats["backend_lines"] += lines
        if path.suffix == ".sql":
            stats["sql_lines"] += lines
        stats["files"].append({"path": relative, "lines": lines})
    _json_write(job_dir / "code_stats.json", stats)


DOC_FONT = "宋体"
DOC_HEADING_FONT = "黑体"
DOC_CODE_FONT = "宋体"
DOC_BODY_SIZE = 12
DOC_TABLE_SIZE = 10.5
DOC_CAPTION_SIZE = 10.5
DOC_HEADER_SIZE = 9
SOURCE_LINE_WIDTH = 72
SOURCE_MARKER_PATTERN = re.compile(r"(?:ai\s*ui\s*enhancer|\bllm\b|prompt\s*:|模型名称)", re.I)


def _set_cell_text_font(run: Any, font_name: str, size: float) -> None:
    run.font.name = font_name
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run._element.rPr.rFonts.set(qn("w:ascii"), font_name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font_name)
    run.font.color.rgb = RGBColor(0, 0, 0)


def _set_style_font(style: Any, font_name: str, size: float) -> None:
    style.font.name = font_name
    style.font.size = Pt(size)
    r_fonts = style.element.get_or_add_rPr().get_or_add_rFonts()
    r_fonts.set(qn("w:eastAsia"), font_name)
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)
    style.font.color.rgb = RGBColor(0, 0, 0)


def _add_field(paragraph: Any, instruction_text: str, fallback: str) -> None:
    """Insert a complex field that Word and WPS both calculate on open."""
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = f" {instruction_text} "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    paragraph.add_run()._r.append(begin)
    paragraph.add_run()._r.append(instruction)
    paragraph.add_run()._r.append(separate)
    placeholder = paragraph.add_run(fallback)
    _set_cell_text_font(placeholder, DOC_FONT, DOC_HEADER_SIZE)
    paragraph.add_run()._r.append(end)


def _add_page_number(paragraph: Any) -> None:
    _add_field(paragraph, "PAGE", "1")


def _enable_field_update(document: Document) -> None:
    settings = document.settings.element
    update = OxmlElement("w:updateFields")
    update.set(qn("w:val"), "true")
    settings.append(update)


def _add_table_of_contents(document: Document) -> None:
    """Insert an update-on-open TOC field; Word/WPS calculates actual page numbers."""
    _enable_field_update(document)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _add_field(paragraph, 'TOC \\o "1-3" \\h \\z \\u', "目录将在打开文档时自动更新")
    for run in paragraph.runs:
        _set_cell_text_font(run, DOC_FONT, DOC_BODY_SIZE)


def _setup_source_doc(document: Document, title: str, version: str) -> None:
    _set_style_font(document.styles["Normal"], DOC_FONT, 10.5)
    document.styles["Normal"].paragraph_format.line_spacing = Pt(12.5)
    section = document.sections[0]
    for margin in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, margin, Cm(2.54))
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_cell_text_font(header.add_run(f"{title} {version} 源代码"), DOC_FONT, 10.5)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_page_number(footer)
    _set_cell_text_font(footer.add_run("/"), DOC_FONT, 10.5)
    _add_field(footer, "NUMPAGES", "1")


def _setup_doc(document: Document, title: str, version: str = "V1.0") -> None:
    section = document.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_cell_text_font(header.add_run(f"{title} {version}".strip()), DOC_FONT, DOC_HEADER_SIZE)
    page = section.header.add_paragraph()
    page.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_cell_text_font(page.add_run("第 "), DOC_FONT, DOC_HEADER_SIZE)
    _add_page_number(page)
    _set_cell_text_font(page.add_run(" 页"), DOC_FONT, DOC_HEADER_SIZE)


def _apply_document_template(document: Document, template: str) -> None:
    style = document.styles["Normal"]
    _set_style_font(style, DOC_FONT, DOC_BODY_SIZE)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(0)
    for level, size in ((1, 16), (2, 14), (3, 12)):
        heading = document.styles[f"Heading {level}"]
        _set_style_font(heading, DOC_HEADING_FONT, size)
        heading.font.bold = True
        heading.paragraph_format.space_before = Pt(16 if level == 1 else 10)
        heading.paragraph_format.space_after = Pt(8 if level == 1 else 6)
        heading.paragraph_format.line_spacing = 1.25


def _source_material_files(root: Path) -> Iterable[Path]:
    """Yield readable business source files, excluding generated/minified presentation assets."""
    for path in _source_files(root):
        relative = path.relative_to(root).as_posix()
        if path.suffix.lower() == ".css" or relative.endswith("/style.css"):
            continue
        yield path


def _source_visual_lines(line: str) -> List[str]:
    if not line:
        return [" "]
    return textwrap.wrap(
        line,
        width=SOURCE_LINE_WIDTH,
        replace_whitespace=False,
        drop_whitespace=False,
        break_long_words=True,
        break_on_hyphens=False,
    ) or [" "]


def generate_source_document(job_dir: Path, job: Optional[Dict[str, Any]] = None) -> None:
    planning = _json_read(job_dir / "planning.json")
    root = job_dir / "generated_project"
    code_lines: List[str] = []
    for path in _source_material_files(root):
        relative = path.relative_to(root).as_posix()
        for line in path.read_text(encoding="utf-8", errors="ignore").splitlines():
            if SOURCE_MARKER_PATTERN.search(line):
                continue
            for visual_line in _source_visual_lines(line):
                code_lines.append(visual_line)
    max_lines = 60 * 50
    if len(code_lines) > max_lines:
        code_lines = code_lines[:1500] + code_lines[-1500:]
    document = Document()
    _setup_source_doc(document, planning["software_name"], (job or {}).get("version", "V1.0"))
    style = document.styles["Normal"]
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = Pt(12.5)
    for index, line in enumerate(code_lines, start=1):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.line_spacing = Pt(12.5)
        run = paragraph.add_run(line)
        _set_cell_text_font(run, DOC_CODE_FONT, 10.5)
    docs = job_dir / "docs"
    docs.mkdir(exist_ok=True)
    document.save(docs / "源代码材料.docx")


def _title(document: Document, text: str, subtitle: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(120)
    _set_cell_text_font(paragraph.add_run(text), DOC_HEADING_FONT, 22)
    sub = document.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_cell_text_font(sub.add_run(subtitle), DOC_FONT, 16)
    document.add_page_break()


def _heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_heading(text, level=level)
    for run in paragraph.runs:
        _set_cell_text_font(run, DOC_HEADING_FONT, {1: 16, 2: 14, 3: 12}[level])


def _body(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(text)
    paragraph.paragraph_format.first_line_indent = Cm(0.74)
    paragraph.paragraph_format.line_spacing = 1.5
    for run in paragraph.runs:
        _set_cell_text_font(run, DOC_FONT, DOC_BODY_SIZE)


def _add_operation_steps(document: Document, steps: List[Dict[str, str]]) -> None:
    """Write actual operations as prose, not as a decorative flow-grid table."""
    for index, step in enumerate(steps, start=1):
        action = step.get("action", "")
        result = step.get("expected_result", "")
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(0.74)
        paragraph.paragraph_format.first_line_indent = Cm(0)
        paragraph.paragraph_format.line_spacing = 1.5
        run = paragraph.add_run(f"第{index}步：{action}。完成后，{result}。")
        _set_cell_text_font(run, DOC_FONT, DOC_BODY_SIZE)


def _add_info_table(document: Document, rows: List[Tuple[str, str]]) -> None:
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        for index, cell in enumerate(cells):
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    _set_cell_text_font(run, DOC_HEADING_FONT if index == 0 else DOC_FONT, DOC_TABLE_SIZE)


def _add_caption(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_cell_text_font(paragraph.add_run(text), DOC_FONT, DOC_CAPTION_SIZE)


def _add_screenshot(document: Document, path: Optional[Path], caption: str, figure: List[int]) -> bool:
    if not path or not path.exists():
        return False
    document.add_picture(str(path), width=Cm(15.5))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_caption(document, f"图 {figure[0]} {caption}")
    figure[0] += 1
    return True


def _create_manual_flow_chart(job_dir: Path) -> Optional[Path]:
    """Create an original, evidence-bounded flow chart for the user manual."""
    try:
        from PIL import Image, ImageDraw, ImageFont
        font_path = Path(os.environ.get("WINDIR", r"C:\Windows")) / "Fonts" / "simsun.ttc"
        font = ImageFont.truetype(str(font_path), 28)
        image = Image.new("RGB", (1800, 250), "white")
        draw = ImageDraw.Draw(image)
        nodes = ["登录系统", "查看首页", "选择功能", "查询或维护", "提交处理", "查看结果"]
        x, y, width, height = 35, 75, 225, 78
        for index, label in enumerate(nodes):
            draw.rounded_rectangle((x, y, x + width, y + height), radius=8, fill="#F2F6FF", outline="#4F6FB8", width=3)
            box = draw.textbbox((0, 0), label, font=font)
            draw.text((x + (width - (box[2] - box[0])) / 2, y + 23), label, fill="#1F2D4D", font=font)
            if index < len(nodes) - 1:
                ax = x + width + 12
                draw.line((ax, y + height / 2, ax + 42, y + height / 2), fill="#333333", width=4)
                draw.polygon([(ax + 42, y + height / 2), (ax + 28, y + height / 2 - 9), (ax + 28, y + height / 2 + 9)], fill="#333333")
            x += 295
        path = job_dir / "manual_flow.png"
        image.save(path)
        return path
    except Exception:
        return None


def _pil_font(size: int, bold: bool = False):
    from PIL import ImageFont

    fonts_dir = Path(os.environ.get("WINDIR", r"C:\Windows")) / "Fonts"
    candidates = [
        fonts_dir / ("simhei.ttf" if bold else "simsun.ttc"),
        fonts_dir / "msyh.ttc",
        fonts_dir / "arial.ttf",
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size)
    return ImageFont.load_default()


def _wrap_text_by_pixel(draw: Any, text: str, font: Any, max_width: int) -> List[str]:
    """Wrap Chinese/English mixed text by rendered pixel width."""
    lines: List[str] = []
    for raw_line in text.split("\n"):
        current = ""
        for char in raw_line:
            candidate = current + char
            bbox = draw.textbbox((0, 0), candidate, font=font)
            if current and bbox[2] - bbox[0] > max_width:
                lines.append(current)
                current = char
            else:
                current = candidate
        if current:
            lines.append(current)
        if not raw_line:
            lines.append("")
    return lines or [text]


def _draw_centered_text(draw: Any, box: Tuple[int, int, int, int], text: str, font: Any, fill: str = "#1F2937") -> None:
    max_width = max(20, box[2] - box[0] - 16)
    lines = _wrap_text_by_pixel(draw, text, font, max_width)
    line_boxes = [draw.textbbox((0, 0), line, font=font) for line in lines]
    line_heights = [(bbox[3] - bbox[1]) for bbox in line_boxes]
    line_gap = 5
    total_height = sum(line_heights) + max(0, len(lines) - 1) * line_gap
    y = max(box[1] + 4, box[1] + (box[3] - box[1] - total_height) / 2)
    for line, bbox, line_height in zip(lines, line_boxes, line_heights):
        width = bbox[2] - bbox[0]
        x = box[0] + (box[2] - box[0] - width) / 2
        if y + line_height > box[3] - 2:
            break
        draw.text((x, y), line, fill=fill, font=font)
        y += line_height + line_gap


def _create_design_architecture_chart(job_dir: Path, planning: Dict[str, Any]) -> Optional[Path]:
    """Create an original layered architecture chart for the design document."""
    try:
        from PIL import Image, ImageDraw

        title_font = _pil_font(30, bold=True)
        layer_font = _pil_font(23, bold=True)
        text_font = _pil_font(20)
        small_font = _pil_font(17)
        image = Image.new("RGB", (1800, 900), "white")
        draw = ImageDraw.Draw(image)
        draw.text((55, 30), f"{planning['software_name']} 分层系统架构", fill="#111827", font=title_font)
        draw.line((55, 78, 1745, 78), fill="#D1D5DB", width=2)

        layers = [
            ("用户访问层", "业务人员、管理人员、统计分析人员\n通过现代浏览器访问"),
            ("前端展示层", "Vue 3、Element Plus、Vite\n登录、首页、菜单、列表、维护表单"),
            ("接口服务层", "REST API、Controller、参数校验\n接收页面请求并返回处理结果"),
            ("业务处理层", "Service、业务规则、状态处理\n模块：" + "、".join(module["name"] for module in planning.get("modules", [])[:5])),
            ("数据持久层", "MyBatis Plus、MySQL\n数据表：" + "、".join(planning.get("database_tables", [])[:5])),
        ]
        colors = ["#EFF6FF", "#ECFDF5", "#FFF7ED", "#F5F3FF", "#F8FAFC"]
        y = 104
        layer_height = 108
        for index, ((label, content), color) in enumerate(zip(layers, colors)):
            draw.rounded_rectangle((80, y, 1720, y + layer_height), radius=18, fill=color, outline="#4B5563", width=3)
            draw.rounded_rectangle((105, y + 22, 330, y + layer_height - 22), radius=12, fill="#FFFFFF", outline="#9CA3AF", width=2)
            _draw_centered_text(draw, (105, y + 22, 330, y + layer_height - 22), label, layer_font)
            _draw_centered_text(draw, (390, y + 12, 1665, y + layer_height - 12), content, text_font)
            if index < len(layers) - 1:
                x = 900
                draw.line((x, y + layer_height + 6, x, y + layer_height + 34), fill="#374151", width=5)
                draw.polygon([(x, y + layer_height + 42), (x - 13, y + layer_height + 22), (x + 13, y + layer_height + 22)], fill="#374151")
            y += layer_height + 48

        draw.text((80, 855), "说明：图中层级、技术栈、模块名称和数据表均来自本次生成的 planning.json 与项目实际生成结果。", fill="#4B5563", font=small_font)
        path = job_dir / "design_architecture.png"
        image.save(path)
        return path
    except Exception:
        return None


def _create_design_process_chart(job_dir: Path, planning: Dict[str, Any]) -> Optional[Path]:
    """Create an original business processing flow chart for the design document."""
    try:
        from PIL import Image, ImageDraw

        title_font = _pil_font(30, bold=True)
        node_font = _pil_font(22, bold=True)
        small_font = _pil_font(17)
        image = Image.new("RGB", (1800, 700), "white")
        draw = ImageDraw.Draw(image)
        draw.text((55, 35), "系统业务处理流程", fill="#111827", font=title_font)
        draw.line((55, 86, 1745, 86), fill="#D1D5DB", width=2)
        nodes = [
            ("登录鉴权", "校验账号并进入首页"),
            ("菜单路由", "选择有权限的功能模块"),
            ("条件查询", "按字段筛选业务记录"),
            ("数据维护", "新增、编辑或执行业务动作"),
            ("接口处理", "REST API 校验并调用服务"),
            ("持久化", "写入或读取 MySQL 数据"),
            ("结果反馈", "刷新列表、详情或提示信息"),
        ]
        width, height = 320, 108
        top_positions = [(115, 160), (530, 160), (945, 160), (1360, 160)]
        bottom_positions = [(320, 360), (735, 360), (1150, 360)]
        positions = top_positions + bottom_positions
        for index, ((label, note), (x, y)) in enumerate(zip(nodes, positions)):
            draw.rounded_rectangle((x, y, x + width, y + height), radius=14, fill="#F2F6FF", outline="#5B6FB0", width=3)
            _draw_centered_text(draw, (x + 12, y + 16, x + width - 12, y + 62), label, node_font, fill="#111827")
            _draw_centered_text(draw, (x + 12, y + 68, x + width - 12, y + height - 12), note, small_font, fill="#374151")
        arrow_pairs = [(0, 1), (1, 2), (2, 3), (3, 4), (4, 5), (5, 6)]
        for start, end in arrow_pairs:
            sx, sy = positions[start]
            ex, ey = positions[end]
            if sy == ey:
                x1, y1 = sx + width + 10, sy + height / 2
                x2, y2 = ex - 10, ey + height / 2
                draw.line((x1, y1, x2, y2), fill="#333333", width=4)
                draw.polygon([(x2, y2), (x2 - 16, y2 - 9), (x2 - 16, y2 + 9)], fill="#333333")
            else:
                x1, y1 = sx + width / 2, sy + height + 10
                x2, y2 = ex + width / 2, ey - 10
                mid_y = (y1 + y2) / 2
                draw.line((x1, y1, x1, mid_y, x2, mid_y, x2, y2), fill="#333333", width=4)
                draw.polygon([(x2, y2), (x2 - 10, y2 - 16), (x2 + 10, y2 - 16)], fill="#333333")
        module_names = "、".join(module["name"] for module in planning.get("modules", [])[:8])
        draw.rounded_rectangle((120, 560, 1680, 645), radius=16, fill="#F8FAFC", outline="#CBD5E1", width=2)
        _draw_centered_text(
            draw,
            (150, 570, 1650, 635),
            f"覆盖功能模块：{module_names}。流程节点只描述实际生成系统的通用处理链路。",
            small_font,
            fill="#374151",
        )
        path = job_dir / "design_process.png"
        image.save(path)
        return path
    except Exception:
        return None


def _add_flow_table(document: Document, steps: List[str]) -> None:
    """Use only for design-level architecture/data-flow diagrams, never manual operations."""
    table = document.add_table(rows=1, cols=len(steps))
    table.style = "Table Grid"
    for index, step in enumerate(steps, start=1):
        cell = table.cell(0, index - 1)
        cell.text = f"{index}\n{step}"
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                _set_cell_text_font(run, DOC_FONT, DOC_CAPTION_SIZE)


def _screenshot_records(job_dir: Path, modules: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    manifest_path = job_dir / "screenshot_manifest.json"
    if manifest_path.exists():
        payload = _json_read(manifest_path)
        records = payload.get("screenshots", [])
        if isinstance(records, list):
            return [item for item in records if isinstance(item, dict)]
    records: List[Dict[str, Any]] = []
    for path in sorted((job_dir / "screenshots").glob("*.png")):
        kind = "other"
        if "login" in path.stem:
            kind = "login"
        elif "dashboard" in path.stem:
            kind = "dashboard"
        for module in modules:
            if module["name"] in path.name:
                kind = "module_create" if "form" in path.stem or "create" in path.stem else "module_list"
                records.append({"kind": kind, "file": path.name, "module_key": module["key"], "module_name": module["name"]})
                break
        else:
            records.append({"kind": kind, "file": path.name})
    return records


def _find_screenshot(job_dir: Path, records: List[Dict[str, Any]], kind: str, module_key: Optional[str] = None) -> Optional[Path]:
    for record in records:
        if record.get("kind") == kind and (module_key is None or record.get("module_key") == module_key):
            path = job_dir / "screenshots" / str(record.get("file", ""))
            if path.exists():
                return path
    return None


def _module_actions(planning: Dict[str, Any], module: Dict[str, Any]) -> List[str]:
    labels = ["查询", "查看详情", "新增", "编辑", "删除", "导出"]
    action_labels = {
        "approve": "通过", "reject": "驳回", "quick_audit": "快速审核", "transfer": "转交",
        "return": "退回补充", "archive": "归档", "submit": "提交", "assign": "分派", "close": "办结",
    }
    pattern = re.compile(rf"^(?:GET|POST|PUT|DELETE|PATCH) /api/{re.escape(module['key'])}/\{{id\}}/([a-zA-Z0-9_-]+)$", re.I)
    for api in planning.get("api_list", []):
        match = pattern.match(api.strip()) if isinstance(api, str) else None
        if match:
            code = match.group(1).replace("-", "_")
            labels.append(action_labels.get(code, code.replace("_", " ").title()))
    return labels


def generate_documents(job: Dict[str, Any], job_dir: Path) -> None:
    planning = _json_read(job_dir / "planning.json")
    stats = _json_read(job_dir / "code_stats.json")
    name = planning["software_name"]
    docs_dir = job_dir / "docs"
    docs_dir.mkdir(exist_ok=True)
    screenshot_records = _screenshot_records(job_dir, planning["modules"])
    document_template = job.get("document_template", "standard")
    actions_by_key = {module["key"]: _module_actions(planning, module) for module in planning["modules"]}
    narratives, narrative_meta = build_document_narratives(planning, actions_by_key)
    _json_write(job_dir / "document_narratives.json", {"meta": narrative_meta, "modules": narratives})

    design = Document()
    _setup_doc(design, name + " 设计说明书", job.get("version", "V1.0"))
    _apply_document_template(design, document_template)
    _title(design, name, "软件设计说明书")
    design_figure = [1]
    _heading(design, "1 项目概述")
    _body(design, planning["description"])
    _body(design, f"本系统采用前后端分离架构，包含 {len(planning['modules'])} 个业务模块，实际源码共 {stats['total_lines']} 行。")
    _body(design, "设计说明书围绕系统目标、总体架构、功能模块、数据接口和界面截图展开，所列模块、字段、接口和截图均以本次生成项目的规划文件及实际产物为依据。")
    _heading(design, "2 总体设计")
    _body(design, "系统采用典型的浏览器/服务端/数据库分层设计。前端采用 Vue 3、Element Plus 与 Vite 组织页面、菜单、表格和表单；后端采用 Java 17、Spring Boot 3 和 MyBatis Plus 提供 REST API 与业务服务；生产数据使用 MySQL 存储。")
    _body(design, "总体数据流为：用户在浏览器界面发起查询、登记、维护或业务动作请求，前端通过 REST API 将请求提交至后端控制器，业务服务完成参数校验、规则处理和数据访问，最终将处理结果返回页面并刷新展示。")
    _add_screenshot(design, _create_design_architecture_chart(job_dir, planning), "分层系统架构图", design_figure)
    _add_screenshot(design, _create_design_process_chart(job_dir, planning), "系统业务处理流程图", design_figure)
    _add_screenshot(design, _find_screenshot(job_dir, screenshot_records, "dashboard"), "系统首页与业务概览", design_figure)
    _heading(design, "3 功能设计")
    _body(design, "功能设计以用户菜单为边界，每个模块保持统一的查询、维护、业务处理和结果反馈模式。模块页面中的字段、按钮和表单来自规划确认结果，接口路径由生成器同步落地到前端 API、后端 Controller 与 Service。")
    for index, module in enumerate(planning["modules"], start=1):
        narrative = narratives[module["key"]]
        _heading(design, f"3.{index} {module['name']}", 2)
        _body(design, f"功能目标：{narrative['overview']}")
        _body(design, f"模块定位：{module.get('description') or module['name']}。该模块通过“{ '、'.join(module['pages']) }”等页面承载业务信息展示和维护入口。")
        _add_info_table(design, [
            ("功能页面", "、".join(module["pages"])),
            ("主要字段", "、".join(module["fields"])),
            ("可用操作", "、".join(actions_by_key[module["key"]])),
            ("处理结果", f"完成{module['name']}数据的查询、登记、维护或业务办理。"),
        ])
        _body(design, "处理流程：进入功能菜单后，根据查询条件定位业务记录；新增或编辑时填写页面字段并提交保存；对已生成的业务动作，系统在确认后调用对应接口并刷新处理结果。")
        _add_flow_table(design, ["进入菜单", "查询记录", "维护/办理", "查看结果"])
        _add_screenshot(design, _find_screenshot(job_dir, screenshot_records, "module_list", module["key"]), f"{module['name']}功能页", design_figure)
        _add_screenshot(design, _find_screenshot(job_dir, screenshot_records, "module_create", module["key"]), f"{module['name']}新增或编辑表单", design_figure)
    _heading(design, "4 数据与接口设计")
    _add_info_table(design, [("数据库表", "、".join(planning["database_tables"])), ("接口数量", str(len(planning["api_list"])))])
    for api in planning["api_list"]:
        _body(design, f"接口：{api}")
    _heading(design, "5 截图索引")
    for record in screenshot_records:
        _body(design, f"{record.get('file', '')}：{record.get('label') or record.get('module_name') or record.get('kind')}")
    design.save(docs_dir / "设计说明书.docx")

    manual = Document()
    _setup_doc(manual, name + " 用户操作手册", job.get("version", "V1.0"))
    _apply_document_template(manual, document_template)
    _title(manual, name, "用户操作手册")
    manual_figure = [1]
    _heading(manual, "目录")
    _add_table_of_contents(manual)
    manual.add_page_break()
    _heading(manual, "一、软件概述")
    _body(manual, planning["description"])
    _body(manual, f"{name}围绕已确认的业务规划，提供" + "、".join(item["name"] for item in planning["modules"]) + "等功能模块，并通过统一菜单和页面操作完成业务数据的查询、维护和办理。")
    _heading(manual, "二、产品背景与目标")
    _body(manual, "本软件面向实际业务资料分散、查询维护效率不足的场景建设。系统目标是在已确认的业务范围内形成统一的功能入口、数据维护界面和处理结果反馈，帮助使用人员按规范完成日常操作。")
    _heading(manual, "三、目标用户")
    _body(manual, "目标用户：" + planning.get("target_users", "系统业务管理人员") + "。不同账号可见菜单以系统实际权限配置为准。")
    _heading(manual, "四、软件特色与优势")
    _body(manual, "系统采用前后端分离方式组织功能页面与业务接口；模块字段、可用操作和截图均以已确认规划及实际生成页面为依据。统一的查询、维护和结果反馈过程有助于减少重复录入并提升操作可追溯性。")
    _heading(manual, "五、使用流程简述")
    _body(manual, "用户登录后进入系统首页，从菜单选择有权限的功能模块；在功能页面中查询或维护业务信息，提交后根据页面提示查看处理结果。以下流程图仅描述系统已生成的通用操作链路。")
    flow_chart = _create_manual_flow_chart(job_dir)
    _add_screenshot(manual, flow_chart, "系统通用使用流程", manual_figure)
    _heading(manual, "六、技术特点")
    _body(manual, "前端采用 Vue 3、Element Plus 与 Vite 构建业务页面，后端采用 Java 17、Spring Boot 3 与 MyBatis Plus 提供 REST 接口，数据通过 MySQL 进行持久化存储。")
    _heading(manual, "七、系统运行环境")
    _add_info_table(manual, [("访问方式", "现代浏览器"), ("运行支撑环境", "Java 17、MySQL"), ("前端技术", "Vue 3、Element Plus、Vite"), ("后端技术", "Spring Boot 3、MyBatis Plus")])
    _heading(manual, "八、登录与首页")
    _heading(manual, "8.1 功能简介", 2)
    _body(manual, "登录页用于完成身份验证并进入系统；首页集中展示业务概览、常用提示和当前用户可访问的功能菜单，是进入各项业务操作的统一入口。")
    _heading(manual, "8.2 使用前提", 2)
    _body(manual, "已使用现代浏览器打开系统访问地址，并已取得有效的登录账号。")
    _heading(manual, "8.3 操作说明", 2)
    _add_operation_steps(manual, [
        {"action": "打开系统登录页，确认浏览器地址正确", "expected_result": "页面显示账号、密码等登录输入区域"},
        {"action": "在登录区域填写已分配的账号和密码", "expected_result": "输入内容可被系统接收，未出现必填校验提示"},
        {"action": "点击登录按钮并等待系统校验", "expected_result": "验证通过后自动进入系统首页"},
        {"action": "在首页查看业务概览，并从侧边或顶部菜单选择需要办理的功能", "expected_result": "页面切换至对应的功能入口"},
    ])
    _heading(manual, "8.4 处理结果与注意事项", 2)
    _body(manual, "登录成功后，系统仅展示当前账号具有权限的菜单。登录失败或页面未跳转时，应先核对账号信息和网络连接，再联系系统管理人员处理。")
    _add_screenshot(manual, _find_screenshot(job_dir, screenshot_records, "login"), "系统登录界面", manual_figure)
    _add_screenshot(manual, _find_screenshot(job_dir, screenshot_records, "dashboard"), "系统首页与功能菜单", manual_figure)
    _heading(manual, "九、功能介绍与操作说明")
    for index, module in enumerate(planning["modules"], start=1):
        narrative = narratives[module["key"]]
        _heading(manual, f"9.{index} {module['name']}", 2)
        _heading(manual, f"9.{index}.1 功能简介", 3)
        _body(manual, narrative["overview"])
        _body(manual, "操作入口：登录系统后，从功能菜单中选择“" + module["name"] + "”进入对应页面。")
        if narrative["preconditions"]:
            _body(manual, "使用前提：" + "；".join(narrative["preconditions"]))
        _heading(manual, f"9.{index}.2 操作说明", 3)
        _add_operation_steps(manual, narrative["steps"])
        _heading(manual, f"9.{index}.3 处理结果与注意事项", 3)
        _body(manual, f"完成操作后，系统会刷新{module['name']}的列表或详情信息；请根据页面提示确认本次操作是否提交成功。")
        if narrative["notes"]:
            _body(manual, "注意事项：" + "；".join(narrative["notes"]))
        _add_screenshot(manual, _find_screenshot(job_dir, screenshot_records, "module_list", module["key"]), f"{module['name']}功能入口与操作区", manual_figure)
        _add_screenshot(manual, _find_screenshot(job_dir, screenshot_records, "module_create", module["key"]), f"{module['name']}数据维护表单", manual_figure)
    manual.save(docs_dir / "用户操作手册.docx")

    application = Document()
    _setup_doc(application, name + " 软件著作权申请信息表", job.get("version", "V1.0"))
    _apply_document_template(application, document_template)
    _title(application, name, "软件著作权申请信息表")
    _heading(application, "1 软件基本信息")
    _add_info_table(
        application,
        [
            ("软件全称", name),
            ("软件简称", name.replace("系统", "")[:20]),
            ("版本号", job.get("version", "V1.0")),
            ("软件分类", planning["software_type"]),
            ("开发完成日期", job.get("completion_date", "")),
            ("发表状态", job.get("publication_status", "未发表")),
            ("著作权人", job.get("applicant_name", "待填写")),
            ("开发方式", "独立开发"),
            ("源程序量", f"{stats['total_lines']} 行"),
        ],
    )
    _heading(application, "2 软件功能与技术特点")
    _body(application, planning["description"])
    _body(
        application,
        "主要功能模块包括：" + "、".join(module["name"] for module in planning["modules"]) + "。",
    )
    _body(
        application,
        "技术架构采用 Vue 3、Element Plus、Java 17、Spring Boot 3、MyBatis Plus 与 MySQL，实现前后端分离和业务数据持久化。",
    )
    _heading(application, "3 开发与运行环境")
    _add_info_table(
        application,
        [
            ("开发硬件环境", job.get("development_hardware", "待申请人确认")),
            ("运行硬件环境", job.get("runtime_hardware", "待申请人确认")),
            ("开发操作系统", job.get("development_os", "Windows 10/11 或兼容环境")),
            ("开发工具", "JDK 17、Maven、Node.js、Vue 3、Python"),
            ("运行平台/操作系统", job.get("runtime_os", "Windows 10/11、Linux 或兼容环境")),
            ("运行支撑环境", "Java 17、MySQL、现代浏览器"),
            ("编程语言", "Java、JavaScript、SQL、Python"),
        ],
    )
    _heading(application, "4 提交前确认")
    _body(application, "本表由系统根据 planning.json 和真实源码统计自动生成。著作权人、版本号、开发完成日期及硬件环境等申请主体信息应在正式提交前人工复核。")
    application.save(docs_dir / "软件著作权申请信息表.docx")


def run_compliance_check(job: Dict[str, Any], job_dir: Path) -> None:
    report = build_compliance_report(job_dir)
    _json_write(job_dir / "compliance_report.json", report)

    document = Document()
    planning = _json_read(job_dir / "planning.json")
    _setup_doc(document, planning["software_name"] + " 合规检查报告", job.get("version", "V1.0"))
    _apply_document_template(document, job.get("document_template", "standard"))
    _title(document, planning["software_name"], "软著材料合规检查报告")
    _heading(document, "1 模拟评分")
    _body(
        document,
        f"综合得分：{report['score']}/{report['max_score']}，评级：{report['grade']}，"
        f"检查结论：{'通过' if report['passed'] else '需整改'}。",
    )
    _heading(document, "2 检查明细")
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["检查项", "结果", "得分", "说明"]
    for cell, value in zip(table.rows[0].cells, headers):
        cell.text = value
    for item in report["items"]:
        cells = table.add_row().cells
        cells[0].text = item["name"]
        cells[1].text = "通过" if item["passed"] else "不通过"
        cells[2].text = f"{item['points']}/{item['max_points']}"
        cells[3].text = item["detail"]
    _heading(document, "3 整改建议")
    if report["suggestions"]:
        for index, suggestion in enumerate(report["suggestions"], start=1):
            _body(document, f"{index}. {suggestion}")
    else:
        _body(document, "未发现需要整改的关键问题。正式提交前仍应人工核对申请主体、版本号和日期。")
    document.save(job_dir / "docs" / "软著材料合规检查报告.docx")
    _update(
        job["job_id"],
        compliance_score=report["score"],
        compliance_grade=report["grade"],
        compliance_passed=report["passed"],
    )


def _zip_dir(source: Path, target: Path, base: Optional[Path] = None) -> None:
    base = base or source
    with zipfile.ZipFile(target, "w", zipfile.ZIP_DEFLATED) as archive:
        for path in sorted(source.rglob("*")):
            if path.is_file() and not any(
                part in {"node_modules", "target", "dist", "__pycache__"}
                for part in path.parts
            ):
                archive.write(path, path.relative_to(base))


def build_package(job_dir: Path) -> None:
    generated_zip = job_dir / "generated_project.zip"
    _zip_dir(job_dir / "generated_project", generated_zip, job_dir / "generated_project")
    readme = job_dir / "README_软著材料说明.md"
    planning = _json_read(job_dir / "planning.json")
    stats = _json_read(job_dir / "code_stats.json")
    _write(
        readme,
        f"""# {planning['software_name']} 软著材料说明

- 生成时间：{_now()}
- 真实源码总行数：{stats['total_lines']}
- 前端代码行数：{stats['frontend_lines']}
- 后端代码行数：{stats['backend_lines']}
- SQL 代码行数：{stats['sql_lines']}

材料由任务规划、固定模板和可选 AI 代码增强自动生成，文档功能范围与 planning.json 保持一致。
截图覆盖清单见 screenshot_manifest.json，可用于核对每个菜单模块的功能页和数据维护表单。
代码增强详情见 enhancement.json。
""",
    )
    package = job_dir / "copyright_package.zip"
    with zipfile.ZipFile(package, "w", zipfile.ZIP_DEFLATED) as archive:
        archive.write(generated_zip, generated_zip.name)
        archive.write(job_dir / "planning.json", "planning.json")
        if (job_dir / "enhancement.json").exists():
            archive.write(job_dir / "enhancement.json", "enhancement.json")
        archive.write(job_dir / "code_stats.json", "code_stats.json")
        if (job_dir / "screenshot_manifest.json").exists():
            archive.write(job_dir / "screenshot_manifest.json", "screenshot_manifest.json")
        if (job_dir / "document_narratives.json").exists():
            archive.write(job_dir / "document_narratives.json", "document_narratives.json")
        if (job_dir / "compliance_report.json").exists():
            archive.write(job_dir / "compliance_report.json", "compliance_report.json")
        archive.write(readme, readme.name)
        for path in sorted((job_dir / "docs").glob("*.docx")):
            archive.write(path, path.name)
        for path in sorted((job_dir / "screenshots").glob("*.png")):
            archive.write(path, f"screenshots/{path.name}")
