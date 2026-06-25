import json
import zipfile
from pathlib import Path
from typing import Any, Dict, List

from docx import Document


def _docx_text(path: Path) -> str:
    document = Document(path)
    parts: List[str] = []
    for section in document.sections:
        parts.extend(paragraph.text for paragraph in section.header.paragraphs)
        parts.extend(paragraph.text for paragraph in section.footer.paragraphs)
    for paragraph in document.paragraphs:
        parts.append(paragraph.text)
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def _docx_visual_count(path: Path) -> int:
    return len(Document(path).inline_shapes)


def _docx_has_page_field(path: Path) -> bool:
    try:
        with zipfile.ZipFile(path) as archive:
            header_xml = "\n".join(
                archive.read(name).decode("utf-8", errors="ignore")
                for name in archive.namelist()
                if (name.startswith("word/header") or name.startswith("word/footer")) and name.endswith(".xml")
            )
    except (OSError, zipfile.BadZipFile):
        return False
    return " PAGE " in header_xml or 'w:instr="PAGE"' in header_xml


def _docx_has_standard_style(path: Path) -> bool:
    try:
        document = Document(path)
        normal = document.styles["Normal"]
        heading = document.styles["Heading 1"]
        font_name = normal.element.get_or_add_rPr().get_or_add_rFonts().get(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia"
        )
        heading_font = heading.element.get_or_add_rPr().get_or_add_rFonts().get(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia"
        )
        if "源代码" in path.name:
            return font_name == "宋体" and normal.font.size is not None
        return font_name == "宋体" and normal.font.size is not None and heading_font == "黑体" and heading.font.size is not None
    except (OSError, ValueError, KeyError):
        return False


SOURCE_FORBIDDEN_MARKERS = ("ai ui enhancer", "llm", "prompt:", "模型名称")


def _screenshot_manifest(job_dir: Path) -> List[Dict[str, Any]]:
    path = job_dir / "screenshot_manifest.json"
    if not path.exists():
        return []
    try:
        payload = json.loads(path.read_text(encoding="utf-8-sig"))
    except (OSError, json.JSONDecodeError):
        return []
    items = payload.get("screenshots", [])
    return [item for item in items if isinstance(item, dict)] if isinstance(items, list) else []


def _check(
    items: List[Dict[str, Any]],
    key: str,
    name: str,
    passed: bool,
    points: int,
    detail: str,
    suggestion: str = "",
) -> None:
    items.append(
        {
            "key": key,
            "name": name,
            "passed": passed,
            "points": points if passed else 0,
            "max_points": points,
            "detail": detail,
            "suggestion": "" if passed else suggestion,
        }
    )


def build_compliance_report(job_dir: Path) -> Dict[str, Any]:
    planning = json.loads((job_dir / "planning.json").read_text(encoding="utf-8"))
    stats = json.loads((job_dir / "code_stats.json").read_text(encoding="utf-8"))
    project = job_dir / "generated_project"
    app_vue = (project / "frontend" / "src" / "App.vue").read_text(
        encoding="utf-8", errors="ignore"
    )
    backend_code = "\n".join(
        path.read_text(encoding="utf-8", errors="ignore")
        for path in (project / "backend" / "src" / "main" / "java").rglob("*.java")
    )
    design_text = _docx_text(job_dir / "docs" / "设计说明书.docx")
    manual_text = _docx_text(job_dir / "docs" / "用户操作手册.docx")
    source_path = job_dir / "docs" / "源代码材料.docx"
    screenshots = list((job_dir / "screenshots").glob("*.png"))
    manifest = _screenshot_manifest(job_dir)
    modules = planning["modules"]
    items: List[Dict[str, Any]] = []

    missing_frontend = [
        module["name"] for module in modules if module["name"] not in app_vue
    ]
    _check(
        items,
        "code_frontend_modules",
        "前端代码与规划模块一致",
        not missing_frontend,
        15,
        "全部模块均存在于前端代码"
        if not missing_frontend
        else "前端缺少：" + "、".join(missing_frontend),
        "补齐规划模块对应的菜单和页面。",
    )

    missing_backend = [
        module["key"]
        for module in modules
        if f'/api/{module["key"]}' not in backend_code
    ]
    _check(
        items,
        "code_backend_modules",
        "后端接口与规划模块一致",
        not missing_backend,
        15,
        "全部模块均有后端 API"
        if not missing_backend
        else "后端缺少：" + "、".join(missing_backend),
        "为缺失模块生成 GET/POST API。",
    )

    screenshot_names = [path.stem for path in screenshots]
    missing_screenshots = [module["name"] for module in modules if not any(module["name"] in name for name in screenshot_names)]
    screenshot_base_ok = any("login" in name for name in screenshot_names) and any("dashboard" in name for name in screenshot_names)
    _check(
        items,
        "screenshot_modules",
        "截图与规划模块一致",
        not missing_screenshots and screenshot_base_ok,
        15,
        f"共 {len(screenshots)} 张截图"
        if not missing_screenshots and screenshot_base_ok
        else "缺少：" + "、".join(missing_screenshots or ["登录页或首页"]),
        "重新运行截图 Agent，确保登录页、首页及全部确认模块均有截图。",
    )

    missing_feature_shots = []
    for module in modules:
        kinds = {item.get("kind") for item in manifest if item.get("module_key") == module["key"]}
        if not {"module_list", "module_create"}.issubset(kinds):
            missing_feature_shots.append(module["name"])
    manifest_base_ok = {item.get("kind") for item in manifest}.issuperset({"login", "dashboard"})
    _check(
        items,
        "screenshot_feature_coverage",
        "截图覆盖菜单功能与数据维护入口",
        bool(manifest) and manifest_base_ok and not missing_feature_shots,
        10,
        f"截图清单覆盖登录、首页及 {len(modules)} 个模块的功能页和表单"
        if manifest and manifest_base_ok and not missing_feature_shots
        else "缺少功能截图清单或模块入口：" + "、".join(missing_feature_shots or ["登录页、首页"]),
        "重新生成截图，确保每个模块都有功能页和新增/编辑表单截图，并写入 screenshot_manifest.json。",
    )

    missing_design = [
        module["name"] for module in modules if module["name"] not in design_text
    ]
    _check(
        items,
        "design_modules",
        "设计说明书覆盖全部模块",
        not missing_design,
        15,
        "设计说明书模块完整"
        if not missing_design
        else "缺少：" + "、".join(missing_design),
        "按 planning.json 补齐功能设计章节。",
    )

    missing_manual = [
        module["name"] for module in modules if module["name"] not in manual_text
    ]
    _check(
        items,
        "manual_modules",
        "用户手册覆盖全部模块",
        not missing_manual,
        15,
        "用户手册模块完整"
        if not missing_manual
        else "缺少：" + "、".join(missing_manual),
        "按 planning.json 补齐操作说明。",
    )

    expected_design_images = 3 + len(modules) * 2
    expected_manual_images = 2 + len(modules) * 2
    design_visuals = _docx_visual_count(job_dir / "docs" / "设计说明书.docx")
    manual_visuals = _docx_visual_count(job_dir / "docs" / "用户操作手册.docx")
    _check(
        items,
        "document_visual_coverage",
        "说明书图文覆盖完整",
        design_visuals >= expected_design_images and manual_visuals >= expected_manual_images,
        10,
        f"设计说明书插图 {design_visuals} 张，用户手册插图 {manual_visuals} 张"
        if design_visuals >= expected_design_images and manual_visuals >= expected_manual_images
        else f"设计说明书插图 {design_visuals}/{expected_design_images}，用户手册插图 {manual_visuals}/{expected_manual_images}",
        "为设计说明书插入架构图、流程图、首页截图和每个模块的功能页/数据维护表单截图，并为用户手册插入登录、首页和模块操作截图。",
    )

    detailed_manual = all(
        f"9.{index}.2 操作说明" in manual_text
        and f"9.{index}.3 处理结果与注意事项" in manual_text
        and module["name"] in manual_text
        for index, module in enumerate(modules, start=1)
    ) and manual_text.count("第1步：") >= len(modules) + 1
    detailed_design = all(
        "功能目标" in design_text and module["name"] in design_text
        for module in modules
    )
    _check(
        items,
        "document_operation_depth",
        "功能说明与操作步骤充分",
        detailed_manual and detailed_design,
        10,
        "全部模块均有功能目标和编号操作步骤"
        if detailed_manual and detailed_design
        else "存在缺少功能目标或编号操作步骤的模块",
        "按模块补充功能目标、字段/API说明和至少四步可执行操作。",
    )

    source_text = _docx_text(source_path).lower() if source_path.exists() else ""
    source_marker_hits = [marker for marker in SOURCE_FORBIDDEN_MARKERS if marker in source_text]
    _check(
        items,
        "source_material",
        "源码材料有效且无生成痕迹",
        source_path.exists()
        and source_path.stat().st_size > 1000
        and stats["total_lines"] > 0
        and not source_marker_hits,
        10,
        f"真实源码 {stats['total_lines']} 行，未发现生成痕迹"
        if not source_marker_hits
        else "源码材料包含禁用标识：" + "、".join(source_marker_hits),
        "重新按材料选材策略生成源码，排除 AI 标记和压缩样式文件。",
    )

    required_docs = {
        "设计说明书.docx",
        "用户操作手册.docx",
        "源代码材料.docx",
        "软件著作权申请信息表.docx",
    }
    current_docs = {path.name for path in (job_dir / "docs").glob("*.docx")}
    missing_docs = sorted(required_docs - current_docs)
    _check(
        items,
        "required_documents",
        "核心申请材料齐全",
        not missing_docs,
        10,
        "核心文档齐全" if not missing_docs else "缺少：" + "、".join(missing_docs),
        "重新执行文档生成阶段。",
    )

    names_consistent = all(
        planning["software_name"] in text
        for text in (design_text, manual_text, _docx_text(source_path))
    )
    _check(
        items,
        "software_name_consistency",
        "软件名称跨材料一致",
        names_consistent,
        5,
        "软件名称一致" if names_consistent else "部分材料未出现标准软件名称",
        "统一 planning.json 与全部文档中的软件名称。",
    )

    layout_docs = [
        job_dir / "docs" / "设计说明书.docx",
        job_dir / "docs" / "用户操作手册.docx",
        source_path,
    ]
    missing_page_fields = [path.name for path in layout_docs if not _docx_has_page_field(path)]
    style_failures = [path.name for path in layout_docs if not _docx_has_standard_style(path)]
    _check(
        items,
        "material_layout",
        "材料页码与统一样式规范",
        not missing_page_fields and not style_failures,
        10,
        "页码域和正文样式符合规范"
        if not missing_page_fields and not style_failures
        else "页码域缺失：" + "、".join(missing_page_fields or ["无"])
        + "；样式不符合：" + "、".join(style_failures or ["无"]),
        "重新生成材料，确保每份 DOCX 使用统一正文样式和 Word/WPS 可识别的 PAGE 域。",
    )

    score = sum(item["points"] for item in items)
    max_score = sum(item["max_points"] for item in items)
    blocking_quality_keys = {
        "screenshot_feature_coverage",
        "document_visual_coverage",
        "document_operation_depth",
        "material_layout",
    }
    blockers = [
        item
        for item in items
        if not item["passed"]
        and (item["max_points"] >= 15 or item["key"] in blocking_quality_keys)
    ]
    if blockers:
        grade = "需整改"
    elif score >= 90:
        grade = "优秀"
    elif score >= 80:
        grade = "良好"
    elif score >= 70:
        grade = "基本合格"
    else:
        grade = "需整改"
    return {
        "score": score,
        "max_score": max_score,
        "grade": grade,
        "passed": score >= 80 and not blockers,
        "items": items,
        "summary": {
            "module_count": len(modules),
            "screenshot_count": len(screenshots),
            "source_lines": stats["total_lines"],
            "document_count": len(current_docs),
        },
        "suggestions": [item["suggestion"] for item in items if item["suggestion"]],
    }
